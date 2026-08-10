"""Focused regression tests for the productization layer."""

from __future__ import annotations

import json
import os
from unittest.mock import patch
import sys
import tempfile
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "mcp-server"))

from core.runtime_paths import ensure_workspace
from core.session_service import compare_session_versions, delete_session
from core.evidence_store import search_official_evidence
from core.standards_catalog import (
    approve_update,
    catalog_paths,
    delete_local_source,
    fetch_update_candidate,
    load_builtin_snapshot,
    load_catalog,
)
from core.visual_qa import find_pdftoppm, inspect_docx_structure, inspect_drawio
from scripts.dc_web import Handler
from scripts.doctor import build_report
from scripts.release_check import run_check


class TestOfficialSnapshotProductContract(unittest.TestCase):
    def test_snapshot_has_traceable_metadata(self):
        snapshot = load_builtin_snapshot()
        self.assertTrue(snapshot.get("snapshot_id"))
        self.assertTrue(snapshot.get("snapshot_version"))
        self.assertGreaterEqual(len(snapshot.get("sources", [])), 8)
        for source in snapshot["sources"]:
            self.assertTrue(source.get("issuer"))
            self.assertTrue(source.get("publication_date"))
            self.assertTrue(source.get("version"))
            self.assertTrue(source.get("grade_levels"))
            self.assertTrue(source.get("retrieved_at"))
            self.assertEqual(source.get("content_hash_status"), "metadata_snapshot_only")
            self.assertTrue(source.get("source_url", "").startswith("https://"))
            self.assertTrue(source.get("clauses"))
            self.assertTrue(source["clauses"][0].get("clause_id"))
            self.assertEqual(source["clauses"][0].get("evidence_status"), "clause_candidate")
            self.assertEqual(source["clauses"][0].get("source_version"), source.get("version"))

    def test_builtin_records_expose_honest_provenance_status(self):
        with tempfile.TemporaryDirectory() as temp:
            catalog = load_catalog(temp)
        for source in catalog["sources"]:
            self.assertEqual(source["content_hash_status"], "metadata_snapshot_only")
            self.assertFalse(source["content_sha256"])
            self.assertEqual(source["provenance"]["verification_status"], "candidate")
            self.assertTrue(source["provenance"]["source_record_sha256"])
            self.assertTrue(source["retrieved_at"])

    def test_search_keeps_source_provenance_on_project_records(self):
        with tempfile.TemporaryDirectory() as temp:
            result = search_official_evidence({"stage": "senior_secondary", "subject": "信息技术", "topic": "Python 循环"}, temp)
        self.assertEqual(result["status"], "found")
        self.assertTrue(result["sources"])
        source = result["sources"][0]
        self.assertTrue(source.get("issuer"))
        self.assertTrue(source.get("grade_levels"))
        self.assertTrue(source.get("stage"))
        self.assertEqual(source.get("content_hash_status"), "metadata_snapshot_only")
        self.assertTrue(source.get("specific_clauses", [])[0].get("source_version"))


class _FakeOfficialResponse:
    def __init__(self, content: bytes, url: str):
        self._content = content
        self._url = url
        self.headers = {"Content-Type": "text/html"}

    def __enter__(self):
        return self

    def __exit__(self, *args):
        return False

    def geturl(self):
        return self._url

    def read(self, size=-1):
        if not self._content:
            return b""
        chunk, self._content = self._content[:size], self._content[size:]
        return chunk


class TestOfficialUpdateHistory(unittest.TestCase):
    def test_update_comparison_and_approval_record_history(self):
        url = "https://www.moe.gov.cn/source-update.html"
        source_id = "local-official-versioned"
        with tempfile.TemporaryDirectory() as temp:
            paths = catalog_paths(temp)
            paths["active"].parent.mkdir(parents=True, exist_ok=True)
            paths["active"].write_text(json.dumps([{
                "source_id": source_id,
                "title": "旧版本",
                "version": "2022",
                "source_url": url,
            }], ensure_ascii=False), encoding="utf-8")
            with patch("core.standards_catalog.urllib.request.urlopen", return_value=_FakeOfficialResponse(b"official bytes", url)):
                staged = fetch_update_candidate(url, temp)
            self.assertEqual(staged["status"], "pending_review")
            self.assertEqual(staged["candidate"]["comparison"]["status"], "same_official_url_pending_version_review")
            update_id = staged["candidate"]["update_id"]
            record = {
                "source_id": source_id,
                "title": "新版本",
                "issuer": "教育主管部门",
                "version": "2025",
                "publication_date": "2025-01-01",
                "effective_date": "2025-09-01",
                "stage": ["senior_secondary"],
                "subject": "信息科技",
                "clauses": [{
                    "clause_id": "NEW-1",
                    "section_path": ["课程目标"],
                    "page_number": "1",
                    "anchor": "课程目标",
                    "excerpt": "原文片段",
                    "normalized_summary": "摘要",
                    "keywords": ["循环"],
                    "applicable_topics": ["Python 循环"],
                    "supports_modules": ["instructional_goal"],
                }],
            }
            approved = approve_update(update_id, temp, teacher_confirmed=True, source_record=record)
            self.assertEqual(approved["status"], "approved")
            self.assertEqual(approved["source"]["content_hash_status"], "retrieved")
            self.assertTrue(approved["source"]["content_sha256"])
            history = json.loads(paths["history"].read_text(encoding="utf-8"))
            self.assertEqual(len(history), 1)
            self.assertEqual(history[0]["status"], "superseded")

    def test_update_approval_requires_complete_source_metadata(self):
        with tempfile.TemporaryDirectory() as temp:
            url = "https://www.moe.gov.cn/source-update.html"
            with patch("core.standards_catalog.urllib.request.urlopen", return_value=_FakeOfficialResponse(b"bytes", url)):
                staged = fetch_update_candidate(url, temp)
            result = approve_update(staged["candidate"]["update_id"], temp, teacher_confirmed=True, source_record={})
        self.assertEqual(result["status"], "blocked")
        self.assertIn("title", result["missing_metadata"])


class TestDrawioPageValidation(unittest.TestCase):
    def test_duplicate_root_ids_are_allowed_across_pages(self):
        xml = """<mxfile><diagram name="A"><mxGraphModel><root><mxCell id="0"/><mxCell id="1" parent="0"/><mxCell id="n1" parent="1"/></root></mxGraphModel></diagram><diagram name="B"><mxGraphModel><root><mxCell id="0"/><mxCell id="1" parent="0"/><mxCell id="n2" parent="1"/></root></mxGraphModel></diagram><diagram name="C"><mxGraphModel><root><mxCell id="0"/><mxCell id="1" parent="0"/><mxCell id="n3" parent="1"/></root></mxGraphModel></diagram></mxfile>"""
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "workbook.drawio"
            path.write_text(xml, encoding="utf-8")
            result = inspect_drawio(path, minimum_pages=3)
        self.assertEqual(result["status"], "pass")
        self.assertEqual(result["duplicate_ids"], [])

    def test_single_page_drawio_is_valid_for_single_page_contract(self):
        xml = """<mxfile><diagram name="single"><mxGraphModel><root><mxCell id="0"/><mxCell id="1" parent="0"/></root></mxGraphModel></diagram></mxfile>"""
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "single.drawio"
            path.write_text(xml, encoding="utf-8")
            result = inspect_drawio(path)
        self.assertEqual(result["status"], "pass")


class TestVisualToolResolution(unittest.TestCase):
    def test_pdftoppm_uses_explicit_environment_path(self):
        with tempfile.TemporaryDirectory() as temp:
            executable = Path(temp) / "pdftoppm.exe"
            executable.write_bytes(b"probe")
            with patch.dict(os.environ, {"PDFTOPPM_PATH": str(executable)}), patch("core.visual_qa.shutil.which", return_value=""):
                self.assertEqual(find_pdftoppm(), str(executable.resolve()))


class TestLocalDeletionBoundaries(unittest.TestCase):
    def test_version_compare_reports_changed_sections(self):
        with tempfile.TemporaryDirectory() as temp:
            dirs = ensure_workspace(temp)
            versions = dirs["projects"] / "session-compare" / "versions"
            (versions / "v001").mkdir(parents=True)
            (versions / "v002").mkdir(parents=True)
            (versions / "v001" / "project.json").write_text(json.dumps({"topic": "循环", "periods": 2}, ensure_ascii=False), encoding="utf-8")
            (versions / "v002" / "project.json").write_text(json.dumps({"topic": "循环", "periods": 3}, ensure_ascii=False), encoding="utf-8")
            result = compare_session_versions("session-compare", 1, 2, temp)
        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["changed_section_count"], 1)
        self.assertEqual(result["changed_sections"][0]["section"], "periods")

    def test_project_delete_stays_inside_workspace(self):
        with tempfile.TemporaryDirectory() as temp:
            dirs = ensure_workspace(temp)
            session_dir = dirs["projects"] / "session-abc"
            session_dir.mkdir(parents=True)
            (session_dir / "session.json").write_text(json.dumps({"session_id": "session-abc"}), encoding="utf-8")
            result = delete_session("session-abc", temp)
            self.assertEqual(result["status"], "deleted")
            self.assertFalse(session_dir.exists())

    def test_builtin_source_cannot_be_deleted(self):
        source_id = load_builtin_snapshot()["sources"][0]["source_id"]
        with tempfile.TemporaryDirectory() as temp:
            result = delete_local_source(source_id, temp)
        self.assertEqual(result["status"], "blocked")

    def test_local_source_can_be_deleted_and_snapshot_rebuilt(self):
        with tempfile.TemporaryDirectory() as temp:
            paths = catalog_paths(temp)
            paths["active"].parent.mkdir(parents=True, exist_ok=True)
            paths["active"].write_text(json.dumps([{"source_id": "local-test-source", "title": "教师确认来源"}], ensure_ascii=False), encoding="utf-8")
            result = delete_local_source("local-test-source", temp)
        self.assertEqual(result["status"], "deleted")
        self.assertTrue(result["snapshot"]["path"])


class TestLocalWebContracts(unittest.TestCase):
    def test_web_api_exposes_health_and_session_project_routes(self):
        class Server:
            workspace = None

        handler = Handler.__new__(Handler)
        handler.server = Server()
        with tempfile.TemporaryDirectory() as temp:
            handler.server.workspace = temp
            health_status, health = handler._api("GET", "health")
            projects_status, projects = handler._api("GET", "projects")
            create_status, created = handler._api(
                "POST",
                "projects",
                {
                    "education_scope": "k12_info_technology",
                    "mode": "standard_fast",
                    "stage": "junior_secondary",
                    "grade": "七年级",
                    "subject": "信息科技",
                    "topic": "算法",
                    "class_profile": {"class_size": 40},
                },
            )
            session_id = created["session"]["session_id"]
            session_status, session = handler._api("GET", f"sessions/{session_id}")
        self.assertEqual(health_status, 200)
        self.assertIn("workspace", health)
        self.assertEqual(projects_status, 200)
        self.assertIn("sessions", projects)
        self.assertEqual(create_status, 200)
        self.assertEqual(session_status, 200)
        self.assertEqual(session["session_id"], session_id)

    def test_web_api_exposes_sources_history_and_private_knowledge(self):
        class Server:
            workspace = None

        handler = Handler.__new__(Handler)
        handler.server = Server()
        with tempfile.TemporaryDirectory() as temp:
            handler.server.workspace = temp
            status, sources = handler._api("GET", "sources")
            knowledge_status, knowledge = handler._api("GET", "knowledge")
        self.assertEqual(status, 200)
        self.assertIn("source_history", sources)
        self.assertEqual(knowledge_status, 200)
        self.assertIn("matches", knowledge)


class TestVisualGateContracts(unittest.TestCase):
    def test_internal_enum_in_word_is_a_blocking_error(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "bad.docx"
            document = Document()
            document.add_paragraph("candidate")
            document.save(path)
            result = inspect_docx_structure(path)
        self.assertEqual(result["status"], "fail")
        self.assertIn("candidate", result["banned_tokens"])

    def test_style_level_chinese_font_counts_as_explicit_font(self):
        from docx import Document
        from docx.shared import Pt

        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "styled.docx"
            document = Document()
            document.styles["Normal"].font.name = "SimSun"
            document.styles["Normal"].font.size = Pt(11)
            document.add_paragraph("教师可读内容")
            document.save(path)
            result = inspect_docx_structure(path)
        self.assertIn("SimSun", result["font_families"])
        self.assertNotIn("未检测到明确的中文字体族", " ".join(result["warnings"]))
        self.assertNotIn("未能从 Word XML 读取显式字号", " ".join(result["warnings"]))


class TestReleaseAndDoctorContracts(unittest.TestCase):
    def test_release_audit_has_no_errors(self):
        self.assertEqual(run_check()["errors"], [])

    def test_doctor_has_no_core_failures(self):
        report = build_report()
        self.assertEqual(report["failure_count"], 0, report)


if __name__ == "__main__":
    unittest.main()
