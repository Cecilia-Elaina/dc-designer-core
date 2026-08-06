"""
Phase 6.2: 标准课程报告级导出重构测试
"""
import sys
import os
import json
import re
import tempfile
import unittest

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'mcp-server'))

SEED_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'examples', 'mvp_algorithm_seed_with_context.json')


def _get_full_project():
    from tools.pipeline import run_mvp_pipeline_with_materials
    with tempfile.TemporaryDirectory() as tmpdir:
        result = run_mvp_pipeline_with_materials(SEED_PATH, tmpdir)
        return result['project']


class TestPhase6Exports(unittest.TestCase):
    """Phase 6.2 标准报告级导出测试"""

    @classmethod
    def setUpClass(cls):
        cls.project = _get_full_project()

    # === Stable filenames ===

    def test_stable_filenames_exist(self):
        from tools.document_exporter import export_all
        with tempfile.TemporaryDirectory() as tmpdir:
            result = export_all(self.project, tmpdir)
            files = result["files"]
            for key in ["dc_report", "lesson_plan", "student_worksheet",
                       "alignment_matrix", "ai_process_record"]:
                self.assertIn(key, files)
                self.assertTrue(os.path.exists(files[key]["path"]))

    def test_export_index_complete(self):
        from tools.document_exporter import export_all
        with tempfile.TemporaryDirectory() as tmpdir:
            result = export_all(self.project, tmpdir)
            with open(result["index_path"], 'r', encoding='utf-8') as f:
                index = json.load(f)
            for field in ["export_status", "warnings", "full_report_docx",
                         "lesson_plan_docx", "student_worksheet_docx",
                         "alignment_matrix_xlsx", "ai_process_record_docx",
                         "source_json", "source_markdown_report",
                         "source_materials_markdown", "files"]:
                self.assertIn(field, index, f"index 缺少 {field}")

    # === Full report structure ===

    def test_full_report_size(self):
        """完整报告文件大小 > 100000 bytes"""
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            export_full_dc_report(self.project, path)
            size = os.path.getsize(path)
            self.assertGreater(size, 100000, f"报告太小: {size} bytes")

    def test_full_report_sections(self):
        """完整报告至少 3 个 section"""
        from docx import Document
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            export_full_dc_report(self.project, path)
            doc = Document(path)
            self.assertGreaterEqual(len(doc.sections), 3,
                                   f"只有 {len(doc.sections)} 个 section，需要至少 3 个")

    def test_full_report_has_header(self):
        """完整报告有页眉"""
        from docx import Document
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            export_full_dc_report(self.project, path)
            doc = Document(path)
            has_header = False
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    for p in section.header.paragraphs:
                        if p.text.strip():
                            has_header = True
                            break
            self.assertTrue(has_header, "报告没有页眉")

    def test_full_report_has_landscape(self):
        """完整报告至少一个横向 section"""
        from docx import Document
        from docx.shared import Emu
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            export_full_dc_report(self.project, path)
            doc = Document(path)
            has_landscape = False
            for section in doc.sections:
                if section.page_width > section.page_height:
                    has_landscape = True
                    break
            self.assertTrue(has_landscape, "报告没有横向页面")

    def test_full_report_table_count(self):
        """完整报告表格数量 >= 30"""
        from docx import Document
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            export_full_dc_report(self.project, path)
            doc = Document(path)
            self.assertGreaterEqual(len(doc.tables), 30,
                                   f"表格只有 {len(doc.tables)} 个，需要至少 30 个")

    def test_full_report_has_images(self):
        """完整报告包含真实图片"""
        import zipfile
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            export_full_dc_report(self.project, path)
            with zipfile.ZipFile(path, 'r') as z:
                media = [f for f in z.namelist() if 'media' in f]
                self.assertGreaterEqual(len(media), 2,
                                       f"图片只有 {len(media)} 张，需要至少 2 张")
                # Check document.xml has w:drawing
                with z.open('word/document.xml') as f:
                    content = f.read().decode('utf-8')
                    drawing_count = content.count('w:drawing')
                    self.assertGreaterEqual(drawing_count, 2,
                                           f"w:drawing 只有 {drawing_count} 个，需要至少 2 个")

    def test_full_report_no_raw_dict(self):
        """完整报告无 Python dict 原文"""
        import re
        from docx import Document
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            export_full_dc_report(self.project, path)
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            dict_matches = re.findall(r"\{'[a-z_]+':", text)
            self.assertEqual(len(dict_matches), 0, f"Found {len(dict_matches)} raw dict outputs")

    def test_full_report_has_three_parts(self):
        """完整报告包含报告一、报告二、报告三"""
        from docx import Document
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            export_full_dc_report(self.project, path)
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("报告一", text)
            self.assertIn("报告二", text)
            self.assertIn("报告三", text)

    def test_full_report_has_key_content(self):
        """完整报告包含关键内容"""
        from docx import Document
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            export_full_dc_report(self.project, path)
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            # Check for key content in 3-part structure
            checks = [
                ("教学分析", "报告一：教学分析"),
                ("教学策略", "报告二：教学策略"),
                ("形成性评价", "报告三：形成性评价"),
                ("目的分析", "目的分析"),
                ("从属技能", "从属技能分析"),
                ("学习者", "学习者分析"),
                ("教学策略", "教学策略"),
                ("绩效目标", "绩效目标"),
                ("入门技能", "入门技能"),
                ("前测", "前测"),
                ("后测", "后测"),
                ("质量门禁", "质量门禁"),
                ("教师确认", "教师确认"),
                ("一对一评价", "一对一评价"),
                ("小组评价", "小组评价"),
            ]
            for keyword, desc in checks:
                self.assertIn(keyword, text, f"报告缺少: {desc}")

    def test_full_report_has_draft_warning(self):
        """完整报告保留待验证草案风险提示"""
        from docx import Document
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            unverified = dict(self.project)
            unverified["quality_check"] = dict(self.project.get("quality_check", {}))
            unverified["quality_check"]["can_export_as_final"] = False
            export_full_dc_report(unverified, path)
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("待验证草案", text)

    def test_full_report_part3_is_plan(self):
        """报告三必须是'形成性评价方案（待实施）'"""
        from docx import Document
        from tools.document_exporter import export_full_dc_report
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            export_full_dc_report(self.project, path)
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            # Part 3 should mention it's a plan, not completed evaluation
            has_plan_indicator = ("待实施" in text or "方案" in text or "实施后" in text)
            self.assertTrue(has_plan_indicator, "报告三应标注为待实施方案")

    # === AI process record ===

    def test_ai_process_record_exists(self):
        from tools.document_exporter import export_all
        with tempfile.TemporaryDirectory() as tmpdir:
            result = export_all(self.project, tmpdir)
            self.assertIn("ai_process_record", result["files"])
            self.assertTrue(os.path.exists(result["files"]["ai_process_record"]["path"]))

    def test_ai_process_record_tables(self):
        """AI 过程记录表至少有 3 张表"""
        from docx import Document
        from tools.document_exporter import export_ai_process_record
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test_ai.docx")
            export_ai_process_record(self.project, path)
            doc = Document(path)
            self.assertGreaterEqual(len(doc.tables), 3)

    def test_ai_process_record_iterations(self):
        """AI 过程记录表迭代日志至少 8 行"""
        from docx import Document
        from tools.document_exporter import export_ai_process_record
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test_ai.docx")
            export_ai_process_record(self.project, path)
            doc = Document(path)
            # Find the iteration log table (first large table)
            max_rows = 0
            for table in doc.tables:
                if len(table.rows) > max_rows:
                    max_rows = len(table.rows)
            self.assertGreaterEqual(max_rows, 8,
                                   f"最大表只有 {max_rows} 行，需要至少 8 行")

    def test_ai_process_record_content(self):
        """AI 过程记录表包含关键内容"""
        from docx import Document
        from tools.document_exporter import export_ai_process_record
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test_ai.docx")
            export_ai_process_record(self.project, path)
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("迭代", text)
            self.assertIn("AI", text)
            has_verify = "验证" in text or "清单" in text or "检查" in text
            self.assertTrue(has_verify)
            has_ethics = "伦理" in text or "诚信" in text or "完整性" in text
            self.assertTrue(has_ethics)

    # === Lesson plan ===

    def test_lesson_plan_sections(self):
        from docx import Document
        from tools.document_exporter import export_lesson_plan
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test_plan.docx")
            export_lesson_plan(self.project, path)
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            for s in ["教学流程", "教学重点", "教学难点"]:
                self.assertIn(s, text)

    # === Student worksheet ===

    def test_student_worksheet_tasks(self):
        from docx import Document
        from tools.document_exporter import export_student_worksheet
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test_ws.docx")
            export_student_worksheet(self.project, path)
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            for t in ["任务一", "任务六", "我的步骤", "自我检查"]:
                self.assertIn(t, text)

    # === Excel ===

    def test_excel_has_data(self):
        from openpyxl import load_workbook
        from tools.document_exporter import export_alignment_matrix
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.xlsx")
            export_alignment_matrix(self.project, path)
            wb = load_workbook(path)
            ws = wb.active
            self.assertGreater(ws.max_row, 1)

    # === Pipeline integration ===

    def test_run_mvp_export_package(self):
        from tools.pipeline import run_mvp_pipeline_with_materials, run_mvp_export_package
        with tempfile.TemporaryDirectory() as tmpdir:
            materials_result = run_mvp_pipeline_with_materials(SEED_PATH, tmpdir)
            project = materials_result['project']
            export_result = run_mvp_export_package(project, tmpdir)
            self.assertIn("files", export_result)
            self.assertIn("index_path", export_result)
            for key in ["dc_report", "lesson_plan", "student_worksheet",
                       "alignment_matrix", "ai_process_record"]:
                self.assertIn(key, export_result["files"])
                self.assertTrue(os.path.exists(export_result["files"][key]["path"]))


if __name__ == "__main__":
    unittest.main()
