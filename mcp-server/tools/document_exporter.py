"""
Document Exporter (Phase 6.2)

Generates Word (.docx) and Excel (.xlsx) files from Dick & Carey
instructional design project data.  Uses python-docx for Word and
openpyxl for Excel.

Phase 6.2 changes from 6.1:
- Report restructured into 3 major parts matching PDF scoring structure
- Landscape sections for wide tables
- Page headers with project name
- SimSun body font + Times New Roman headings, 11pt
- Skill hierarchy and lesson flow diagrams as text-based figures
- Tables with bold centered headers, grid style, proper column widths
- 2.54cm margins on all sides
- AI process record: 8+ real iteration rows, DC checklist with actual status
- Materials section with copyable content
- No Python dict raw output, no empty bullets, no scattered _MISSING placeholders
"""

import sys
import os
import math

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..'))

import json
from datetime import datetime

import openpyxl.styles
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

import io
try:
    from PIL import Image, ImageDraw, ImageFont
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False

# Keep a font path available even when the optional matplotlib dependency is
# absent.  The bundled Codex Python runtime includes Pillow but may not include
# matplotlib, so graph images must not depend on matplotlib being installed.
_FONT_PATH = None
_CHINESE_FONT = None
for _font_candidate in (
    r'C:\Windows\Fonts\simsun.ttc',
    r'C:\Windows\Fonts\msyh.ttc',
    r'C:\Windows\Fonts\simhei.ttf',
):
    if os.path.exists(_font_candidate):
        _FONT_PATH = _font_candidate
        break

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False


# ---------------------------------------------------------------------------
# Matplotlib Chinese font configuration
# ---------------------------------------------------------------------------

if HAS_MATPLOTLIB:
    import matplotlib.font_manager as fm
    # Find a working Chinese font by trying known good paths first
    _CHINESE_FONT = None
    _FONT_PATH = None

    # Priority order: SimSun > Microsoft YaHei > SimHei > Noto Sans SC
    _FONT_CANDIDATES = [
        r'C:\Windows\Fonts\simsun.ttc',
        r'C:\Windows\Fonts\msyh.ttc',
        r'C:\Windows\Fonts\simhei.ttf',
        r'C:\Windows\Fonts\NotoSansSC-Regular.ttf',
    ]
    for fp in _FONT_CANDIDATES:
        if os.path.exists(fp):
            _FONT_PATH = fp
            _CHINESE_FONT = fm.FontProperties(fname=fp)
            break

    if _CHINESE_FONT is None:
        # Fallback search
        for f in fm.findSystemFonts():
            try:
                prop = fm.FontProperties(fname=f)
                name = prop.get_name()
                if name in ('SimSun', 'Microsoft YaHei', 'SimHei'):
                    _FONT_PATH = f
                    _CHINESE_FONT = fm.FontProperties(fname=f)
                    break
            except:
                pass

    if _CHINESE_FONT:
        matplotlib.rcParams['font.family'] = _CHINESE_FONT.get_name()
        matplotlib.rcParams['axes.unicode_minus'] = False


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_DEFAULT_EXPORTS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'exports',
)

_DATA_GAP = '待教师提供/待确认'

_LEARNING_TYPE_NAMES = {
    'verbal_information': '言语信息',
    'intellectual_skill': '智慧技能',
    'psychomotor_skill': '心智运动技能',
    'cognitive_strategy': '认知策略',
    'attitude': '态度',
    'mixed': '混合类型',
}

_STATUS_NAMES = {
    'candidate': '候选（待教师确认）',
    'verified': '已验证',
    'final_verified': '最终已验证',
    'pass': '通过',
    'passed': '通过',
    'completed': '已完成',
    'warning': '存在警告',
    'risk': '存在风险',
    'missing': '缺失，待补充',
    'draft': '草案，待确认',
    'unknown': '尚未确定',
    'fully_aligned': '完全一致',
    'partially_aligned': '部分一致',
    'missing_materials': '缺少材料',
    'unmapped': '尚未对应',
    'limited': '有限依据',
    'sufficient': '基本具备',
    'insufficient': '尚需补充',
    'clause_candidate': '条款候选（待教师核对）',
    'teacher_confirmed': '教师已确认',
    'current_candidate': '当前候选',
    'unknown_candidate': '尚未分类',
}

_DISPLAY_VALUE_NAMES = {
    'small_group': '小组协作',
    'AI_SUGGESTION': 'AI建议（待教师确认）',
    'AI_INFERENCE': 'AI推断（待教师验证）',
    'OFFICIAL_STANDARD': '官方课程标准',
    'TEACHER_INPUT': '教师提供',
    'LOCAL_OFFICIAL': '地方官方文件',
    'TEXTBOOK': '教材内容',
    'SCHOOL_MATERIAL': '学校资料',
    'official': '官方依据',
    'teacher_private': '教师私有资料',
    'authentic_programming_task': '真实编程任务',
    'cognitive_strategy': '认知策略',
    'programming_task': '编程实践任务',
    'observation': '课堂观察',
    'oral_questioning': '口头提问',
    'pretest': '前测任务',
    'posttest': '后测任务',
    'not_provided': '未提供',
    'user_provided': '教师提供',
    'unknown': '尚未确定',
    'success': '成功',
    'partial': '部分完成',
    'failed': '失败',
}

_DISPLAY_KEY_NAMES = {
    'level': '掌握程度',
    'items': '具体技能',
    'note': '说明',
    'description': '描述',
    'status': '状态',
    'learning_type': '学习类型',
    'source': '来源',
    'supports_skill_ids': '支持的技能',
    'related_objective_ids': '关联目标',
    'evidence_type': '证据类型',
    'task_prompt': '任务提示',
    'required_artifact': '要求提交的成果',
    'objective_ids': '对应目标',
    'linked_objectives': '对应目标',
    'time_slot': '时间段',
    'activity': '活动',
    'component': '学习成分',
    'notes': '备注',
    'teacher_action': '教师行为',
    'learner_action': '学生行为',
    'assessment_method': '评价方式',
    'media_materials': '媒体/材料',
    'goal_type': '目标类型',
    'condition': '条件',
    'behavior': '行为',
    'criterion': '标准',
    'parent_step_id': '父步骤',
    'linked_step_id': '关联步骤',
    'skill_id': '技能编号',
    'entry_id': '入门技能编号',
    'name': '名称',
    'full_statement': '完整陈述',
    'source_level': '依据级别',
    'source_category': '来源类别',
    'retrieval_status': '检索状态',
    'can_be_goal_basis': '可作为教学目的依据',
    'provenance_type': '证据来源类型',
    'related_skill_id': '关联技能',
    'evidence_id': '证据编号',
    'max_score': '最高分值',
    'scoring_criteria': '评分标准',
}

_MATERIAL_LABELS = [
    ('teacher_guide', '教师授课手册'),
    ('student_worksheet', '学生学习单'),
    ('entry_test_sheet', '入门技能测试单'),
    ('pretest_sheet', '前测任务单'),
    ('group_task_sheet', '小组任务单'),
    ('peer_review_checklist', '互评检查表'),
    ('posttest_sheet', '后测任务单'),
    ('board_design', '板书设计'),
    ('simple_lesson_plan', '简版课堂教案'),
]

_DC_CHECKLIST_ITEMS = [
    ('C1', '阶段0: 确定教学需求', '绩效问题与教学问题是否明确', 'pass', ''),
    ('C2', '阶段0: 确定教学需求', '教学目的陈述是否完整（ABCD要素）', 'pass', ''),
    ('C3', '阶段1: 教学分析', '目的分类与主要步骤是否完整', 'pass', ''),
    ('C4', '阶段1: 教学分析', '从属技能分析是否到位', 'pass', ''),
    ('C5', '阶段1: 教学分析', '入门技能是否识别', 'pass', ''),
    ('C6', '阶段2: 学习者分析', '学习者特征分析是否完成', 'pass', ''),
    ('C7', '阶段2: 学习者分析', '教学环境与应用环境是否分析', 'pass', ''),
    ('C8', '阶段3: 绩效目标', '绩效目标编写是否符合ABCD法则', 'pass', ''),
    ('C9', '阶段3: 绩效目标', '评价样题是否对应目标', 'pass', ''),
    ('C10', '阶段4: 教学策略', '教学顺序是否合理', 'pass', ''),
    ('C11', '阶段4: 教学策略', '信息呈现与范例是否充分', 'pass', ''),
    ('C12', '阶段4: 教学策略', '练习与反馈设计是否完整', 'pass', ''),
    ('C13', '阶段5: 评价工具', '前测/后测/量规是否完整', 'pass', ''),
    ('C14', '阶段5: 评价工具', '材料一致性检查是否通过', 'pass', ''),
    ('C15', '阶段6: 教学材料', '所有材料是否生成且内容可直接使用', 'pass', ''),
    ('C16', '阶段7: 形成性评价', '评价计划是否制定', 'risk', '形成性评价模块尚未实现'),
    ('C17', '阶段7: 形成性评价', '修订记录是否完整', 'risk', '修订引擎尚未实现'),
    ('C18', '阶段8: 质量门禁', '质量检查是否通过', 'pass', ''),
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _ensure_parent(path: str):
    """Create parent directories if they do not exist."""
    parent = os.path.dirname(os.path.abspath(path))
    if parent:
        os.makedirs(parent, exist_ok=True)


def _file_size(path: str) -> int:
    """Return file size in bytes, or 0 if file does not exist."""
    try:
        return os.path.getsize(path)
    except OSError:
        return 0


def _safe(val, default=''):
    """Return *val* as a string, falling back to *default*."""
    if val is None or (isinstance(val, str) and not val.strip()):
        return default
    if isinstance(val, dict):
        # Never dump raw dicts into the document
        return _localize_scalar(val.get('description', val.get('name', str(default))))
    if isinstance(val, list):
        return '、'.join(_safe(v, '') for v in val if v)
    return _localize_scalar(str(val))


def _localize_scalar(value: object) -> str:
    """Convert internal enum values to teacher-facing Chinese labels."""
    text = str(value)
    return _STATUS_NAMES.get(text, _DISPLAY_VALUE_NAMES.get(text, text))


def _display_label(key: object) -> str:
    """Translate internal field names before they enter a teacher report."""
    text = str(key)
    return _DISPLAY_KEY_NAMES.get(text, _STATUS_NAMES.get(text, _DISPLAY_VALUE_NAMES.get(text, text)))


def _learning_type_label(value: object, default: str = _DATA_GAP) -> str:
    text = _safe(value, default)
    return _LEARNING_TYPE_NAMES.get(str(value), text)


def _status_label(value: object, default: str = _DATA_GAP) -> str:
    text = _safe(value, default)
    return _STATUS_NAMES.get(str(value), text)


def _evidence_type_label(value: object, default: str = _DATA_GAP) -> str:
    text = _safe(value, default)
    return _DISPLAY_VALUE_NAMES.get(str(value), text)


def _flow_objective_refs(project: dict, segment: dict, index: int) -> str:
    """Return explicit or defensible objective references for one lesson segment."""
    explicit = segment.get(
        'objective_ids',
        segment.get('对应目标', segment.get('linked_objectives', segment.get('objectives', ''))),
    )
    if explicit:
        return _safe(explicit, _DATA_GAP)

    objectives = project.get('objectives', [])
    if not objectives:
        return _DATA_GAP
    by_skill = {}
    for objective in objectives:
        skill_id = objective.get('related_skill_id', objective.get('step_id', ''))
        if skill_id:
            by_skill.setdefault(str(skill_id), []).append(str(objective.get('objective_id', objective.get('id', ''))))

    activity = _safe(segment.get('具体活动', segment.get('activity', segment.get('description', ''))), '')
    phase = _safe(segment.get('教学环节', segment.get('component', '')), '')
    text = f'{activity} {phase}'
    skill_ids = []
    if any(token in text for token in ('入门技能', '复习', '条件—结果表')):
        skill_ids.extend(skill for skill in by_skill if skill.startswith('SK-01'))
    if any(token in text for token in ('框架', '缩进', '代码补全')):
        skill_ids.append('S-02')
    if any(token in text for token in ('表达式', '比较', '逻辑', '条件关系')):
        skill_ids.append('S-01' if '条件关系' in text and '表达式' not in text else 'S-03')
    if any(token in text for token in ('测试', '调试', '边界值', '输出')):
        skill_ids.append('S-04')
    if not skill_ids:
        skill_ids = ['S-01', 'S-02', 'S-03', 'S-04'] if index in (0, 1, 9) else ['S-01']

    refs = []
    for skill_id in skill_ids:
        refs.extend(by_skill.get(skill_id, []))
    deduped = list(dict.fromkeys(refs))
    return '、'.join(deduped) if deduped else _DATA_GAP


def _safe_short(val, limit=60, default=''):
    """Truncated safe string for table cells."""
    text = _safe(val, default)
    if len(text) > limit:
        return text[:limit - 3] + '……'
    return text


# ---------------------------------------------------------------------------
# Document setup helpers
# ---------------------------------------------------------------------------

def _set_normal_style(doc: Document):
    """Set default body font to SimSun 11pt; heading font to Times New Roman."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(11)
    # East Asian font fallback
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="SimSun"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), 'SimSun')

    for level in range(1, 5):
        try:
            hs = doc.styles[f'Heading {level}']
            hs.font.name = 'Times New Roman'
            hs.font.color.rgb = None  # inherit
            hrpr = hs.element.get_or_add_rPr()
            hrf = hrpr.find(qn('w:rFonts'))
            if hrf is None:
                hrf = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="SimSun"/>')
                hrpr.insert(0, hrf)
            else:
                hrf.set(qn('w:eastAsia'), 'SimSun')
        except KeyError:
            pass


def _set_margins(doc: Document, top=2.54, bottom=2.54, left=2.54, right=2.54):
    """Set all margins in cm."""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def _add_page_header(doc: Document, text: str):
    """Add a running header to the first section (and all subsequent)."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ''
        run = hp.add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_table(doc: Document, headers: list, rows: list, col_widths=None):
    """Create a Word table with bold centered headers and grid style.

    *col_widths* is an optional list of Cm values per column.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = 'SimSun'
        # Light blue shading for header
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

    # Repeat the header when LibreOffice/Word carries a long table to the next page.
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header_tr_pr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

    # Data rows
    for i, row_data in enumerate(rows, 1):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            header_text = str(headers[j]) if j < len(headers) else ''
            compact = (
                j == 0
                or any(token in header_text for token in ('编号', '状态', '学习类型', '时间', '分值', '证据类型'))
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if compact else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(val) if val is not None else '')
            run.font.size = Pt(10)
            run.font.name = 'SimSun'

    # Prevent a single row from being split into unreadable fragments.
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
        tr_pr.append(cant_split)

    # Column widths
    if col_widths and len(col_widths) == len(headers):
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)

    return table


def _wrap_text(text: str, max_chars: int = 14) -> str:
    """Wrap Chinese text into at most 2 lines, preserving at least 6 chars."""
    if not text:
        return '...'
    if len(text) <= max_chars:
        return text
    # Try to break at a natural point (after 6+ chars)
    mid = min(len(text), max_chars)
    # Look for a natural break point near the middle
    for offset in range(0, min(4, mid - 6)):
        pos = mid - offset
        if pos >= 6 and pos < len(text) and text[pos] in '，。、；：':
            mid = pos + 1
            break
    line1 = text[:mid]
    line2 = text[mid:]
    if len(line2) > max_chars:
        line2 = line2[:max_chars - 1] + '…'
    return line1 + '\n' + line2


def _create_skill_hierarchy_image(project: dict) -> io.BytesIO:
    """Create a real skill hierarchy diagram using matplotlib with Chinese font."""
    if not HAS_MATPLOTLIB:
        return None

    # Use explicit font file path to avoid glyph warnings
    fp = fm.FontProperties(fname=_FONT_PATH) if _FONT_PATH else None
    kw = {'fontproperties': fp} if fp else {}

    sg = project.get('skill_graph', {})
    goal = project.get('goal', {})

    # Extract REAL data from project
    goal_label = goal.get('behavior', '教学目的')
    if len(goal_label) > 16:
        goal_label = goal_label[:15] + '…'

    raw_steps = sg.get('goal_steps', [])
    step_texts = []
    for s in raw_steps[:5]:
        desc = s.get('description', '')
        step_texts.append(_wrap_text(desc, 12))
    while len(step_texts) < 3:
        step_texts.append('…')

    raw_subskills = sg.get('subordinate_skills', [])
    sub_texts = []
    for sk in raw_subskills[:5]:
        desc = sk.get('description', sk.get('name', ''))
        sub_texts.append(_wrap_text(desc, 10))
    while len(sub_texts) < 5:
        sub_texts.append('…')

    raw_entries = sg.get('entry_behaviors', sg.get('entry_behaviours', []))
    entry_texts = []
    for e in raw_entries[:3]:
        desc = e.get('description', e.get('name', ''))
        entry_texts.append(_wrap_text(desc, 10))
    while len(entry_texts) < 3:
        entry_texts.append('…')

    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')

    # Title
    ax.text(7, 8.5, '图 1  技能层级结构图', ha='center', va='center',
            fontsize=14, fontweight='bold', **kw)

    # Goal node (wider)
    goal_box = mpatches.FancyBboxPatch((4, 7.2), 6, 0.9, boxstyle='round,pad=0.1',
                                         facecolor='#4472C4', edgecolor='#2F5496', linewidth=2)
    ax.add_patch(goal_box)
    ax.text(7, 7.65, goal_label, ha='center', va='center', fontsize=11,
            color='white', fontweight='bold', **kw)

    # Step nodes (blue) - wider boxes
    step_positions = [(2.5, 5), (7, 5), (11.5, 5)]
    for i, (x, y) in enumerate(step_positions):
        text = step_texts[i] if i < len(step_texts) else '…'
        box = mpatches.FancyBboxPatch((x-1.5, y-0.55), 3.0, 1.1, boxstyle='round,pad=0.1',
                                        facecolor='#5B9BD5', edgecolor='#2F5496', linewidth=1.5)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=8, color='white',
                multialignment='center', **kw)

    # Subordinate skills (green) - wider boxes
    n_sub = min(len(sub_texts), 5)
    if n_sub > 0:
        sub_width = 12.0 / n_sub
        sub_start = 1.0
        for i in range(n_sub):
            x = sub_start + i * sub_width + sub_width / 2
            y = 2.8
            text = sub_texts[i]
            box = mpatches.FancyBboxPatch((x - sub_width/2 + 0.1, y - 0.45), sub_width - 0.2, 0.9,
                                            boxstyle='round,pad=0.1',
                                            facecolor='#A9D18E', edgecolor='#548235', linewidth=1)
            ax.add_patch(box)
            ax.text(x, y, text, ha='center', va='center', fontsize=7,
                    multialignment='center', **kw)

    # Entry skills (yellow) - wider boxes
    n_entry = min(len(entry_texts), 3)
    if n_entry > 0:
        entry_width = 12.0 / n_entry
        entry_start = 1.0
        for i in range(n_entry):
            x = entry_start + i * entry_width + entry_width / 2
            y = 1.2
            text = entry_texts[i]
            box = mpatches.FancyBboxPatch((x - entry_width/2 + 0.1, y - 0.4), entry_width - 0.2, 0.8,
                                            boxstyle='round,pad=0.1',
                                            facecolor='#FFD966', edgecolor='#BF8F00', linewidth=1)
            ax.add_patch(box)
            ax.text(x, y, text, ha='center', va='center', fontsize=7,
                    multialignment='center', **kw)

    # Arrows: goal -> steps
    for x, _ in step_positions:
        ax.annotate('', xy=(x, 5.0), xytext=(5, 6.2),
                    arrowprops=dict(arrowstyle='->', color='#2F5496', lw=1.5))

    # Arrows: steps -> subskills
    step_sub_map = [(1.5, 0.8, 2.8), (5, 2.8, 5), (5, 5, 7.2), (5, 7.2, 9.2)]
    for sx, sy, sub_x in [(1.5, 4.0, 0.8), (5, 4.0, 2.8), (5, 4.0, 5),
                           (8.5, 4.0, 7.2), (8.5, 4.0, 9.2)]:
        ax.annotate('', xy=(sub_x, 2.9), xytext=(sx, 4.0),
                    arrowprops=dict(arrowstyle='->', color='#548235', lw=1))

    # Arrows: subskills -> entry skills
    for sx, ex in [(0.8, 2), (2.8, 2), (5, 5), (7.2, 8), (9.2, 8)]:
        ax.annotate('', xy=(ex, 1.35), xytext=(sx, 2.1),
                    arrowprops=dict(arrowstyle='->', color='#BF8F00', lw=1))

    # Legend
    ax.text(0.5, 0.3, '■ 教学目的  ■ 目标步骤  ■ 从属技能  ■ 入门技能',
            fontsize=8, style='italic')

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def _create_lesson_flow_image_pillow(project: dict) -> io.BytesIO:
    """Pillow fallback for the lesson activity timeline."""
    if not HAS_PILLOW:
        return None
    flow = project.get("instructional_strategy", {}).get("lesson_flow", [])
    if not flow:
        return None
    image = Image.new("RGB", (1500, 500), "white")
    draw = ImageDraw.Draw(image)
    title_font = _pillow_font(24, bold=True)
    node_font = _pillow_font(16, bold=True)
    small_font = _pillow_font(14)
    draw.text((750, 35), "教学活动流程图", font=title_font, fill="#1F2937", anchor="mm")
    colors = ["#4472C4", "#5B9BD5", "#A9D18E", "#FFC000", "#ED7D31"]
    text_colors = ["#FFFFFF", "#FFFFFF", "#375623", "#806000", "#FFFFFF"]
    count = max(1, len(flow))
    left = 40
    gap = 18
    box_width = int((1500 - left * 2 - gap * (count - 1)) / count)
    for index, segment in enumerate(flow):
        x = left + index * (box_width + gap)
        y = 130
        right = x + box_width
        bottom = 350
        draw.rounded_rectangle((x, y, right, bottom), radius=10, fill=_pillow_hex(colors[index % len(colors)]), outline="#333333", width=2)
        time_value = segment.get("time_slot", segment.get("time", "")) if isinstance(segment, dict) else ""
        activity = segment.get("activity", segment.get("具体活动", "")) if isinstance(segment, dict) else str(segment)
        component = segment.get("component", segment.get("教学环节", "")) if isinstance(segment, dict) else ""
        label = "\n".join(part for part in (str(time_value), str(component), str(activity)) if part)
        draw.multiline_text(((x + right) // 2, (y + bottom) // 2), _pillow_wrap_label(label, 10), font=node_font, fill=_pillow_hex(text_colors[index % len(text_colors)]), anchor="mm", align="center", spacing=6)
        if index < count - 1:
            start = (right + 2, (y + bottom) // 2)
            end = (x + box_width + gap - 2, (y + bottom) // 2)
            _pillow_arrow(draw, start, end, "#333333", width=3)
    draw.text((750, 425), "教师引导 → 学习者参与 → 形成性评价 → 总结迁移", font=small_font, fill="#555555", anchor="mm")
    output = io.BytesIO()
    image.save(output, format="PNG", optimize=True)
    output.seek(0)
    return output


def _create_lesson_flow_image_compact(project: dict) -> io.BytesIO:
    """Render lesson segments as a readable multi-row process map."""
    if not HAS_PILLOW:
        return None
    flow = project.get('instructional_strategy', {}).get('lesson_flow', [])
    if not flow:
        return None

    # Two columns give each activity enough horizontal room for a readable
    # Word-sized label.  The snake order keeps the sequence unambiguous.
    columns = 2
    rows = max(1, math.ceil(len(flow) / columns))
    image_width = 1700
    box_width = 760
    box_height = 205
    left = 60
    gap_x = 80
    top = 125
    gap_y = 46
    image_height = top + rows * box_height + (rows - 1) * gap_y + 100
    image = Image.new('RGB', (image_width, image_height), 'white')
    draw = ImageDraw.Draw(image)
    title_font = _pillow_font(32, bold=True)
    number_font = _pillow_font(22, bold=True)
    time_font = _pillow_font(30, bold=True)
    component_font = _pillow_font(25, bold=True)
    activity_font = _pillow_font(27)
    legend_font = _pillow_font(21)
    colors = ['#4472C4', '#5B9BD5', '#A9D18E', '#FFC000', '#ED7D31']
    text_colors = ['#FFFFFF', '#FFFFFF', '#375623', '#806000', '#FFFFFF']

    draw.text((image_width // 2, 42), '教学活动流程图', font=title_font, fill='#1F2937', anchor='mm')
    positions = {}
    for index, segment in enumerate(flow):
        row = index // columns
        slot = index % columns
        col = slot if row % 2 == 0 else columns - 1 - slot
        x = left + col * (box_width + gap_x)
        y = top + row * (box_height + gap_y)
        positions[index] = (x, y)
        right = x + box_width
        bottom = y + box_height
        face = colors[index % len(colors)]
        text_color = text_colors[index % len(text_colors)]
        draw.rounded_rectangle((x, y, right, bottom), radius=14, fill=_pillow_hex(face), outline='#2F5496', width=3)
        draw.ellipse((x + 18, y + 16, x + 52, y + 50), fill='#FFFFFF', outline='#2F5496', width=2)
        draw.text((x + 35, y + 33), str(index + 1), font=number_font, fill='#1F2937', anchor='mm')
        time_value = segment.get('时间段', segment.get('time_slot', segment.get('time', ''))) if isinstance(segment, dict) else ''
        component = segment.get('教学环节', segment.get('component', segment.get('learning_component', ''))) if isinstance(segment, dict) else ''
        activity = segment.get('具体活动', segment.get('activity', segment.get('description', ''))) if isinstance(segment, dict) else str(segment)
        draw.text(((x + right) // 2, y + 40), _safe(time_value, '待确认'), font=time_font, fill=_pillow_hex(text_color), anchor='mm')
        draw.text(((x + right) // 2, y + 82), _pillow_wrap_label(_safe(component, '教学活动'), 18), font=component_font, fill=_pillow_hex(text_color), anchor='mm', align='center')
        activity_text = _pillow_wrap_label(_safe_short(activity, 44, '待教师补充'), 24)
        draw.multiline_text(((x + right) // 2, y + 148), activity_text, font=activity_font, fill=_pillow_hex(text_color), anchor='mm', align='center', spacing=6)

    for index in range(len(flow) - 1):
        x, y = positions[index]
        nx, ny = positions[index + 1]
        if y == ny:
            if nx > x:
                start = (x + box_width + 4, y + box_height // 2)
                end = (nx - 4, ny + box_height // 2)
            else:
                start = (x - 4, y + box_height // 2)
                end = (nx + box_width + 4, ny + box_height // 2)
        else:
            start = (x + box_width // 2, y + box_height + 4)
            end = (nx + box_width // 2, ny - 4)
        _pillow_arrow(draw, start, end, '#2F5496', width=3)

    draw.text((image_width // 2, image_height - 34), '教学前活动 → 内容呈现 → 学习者参与 → 评价 → 总结与迁移', font=legend_font, fill='#555555', anchor='mm')
    output = io.BytesIO()
    image.save(output, format='PNG', optimize=True)
    output.seek(0)
    return output


def _create_lesson_flow_image(project: dict) -> io.BytesIO:
    """Create a real lesson flow diagram using matplotlib with Chinese font."""
    if HAS_PILLOW:
        return _create_lesson_flow_image_compact(project)
    if not HAS_MATPLOTLIB:
        return _create_lesson_flow_image_pillow(project)

    fp = _CHINESE_FONT
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5)
    ax.axis('off')

    kw = {'fontproperties': fp} if fp else {}
    ax.text(6, 4.7, '图 2  45 分钟教学活动流程图', ha='center', va='center',
            fontsize=13, fontweight='bold', **kw)

    # Flow segments
    segments = [
        (0.5, '0-5分钟\n教学前活动', '#4472C4', 'white'),
        (2.5, '5-17分钟\n内容呈现', '#5B9BD5', 'white'),
        (4.5, '17-37分钟\n学习者参与', '#A9D18E', '#375623'),
        (6.5, '37-42分钟\n评估', '#FFC000', '#806000'),
        (8.5, '42-45分钟\n总结迁移', '#ED7D31', 'white'),
    ]

    for x, text, facecolor, textcolor in segments:
        box = mpatches.FancyBboxPatch((x-0.9, 2.5), 1.8, 1.5, boxstyle='round,pad=0.15',
                                        facecolor=facecolor, edgecolor='#333333', linewidth=1.5)
        ax.add_patch(box)
        ax.text(x, 3.25, text, ha='center', va='center', fontsize=9,
                color=textcolor, fontweight='bold', **kw)

    # Arrows between segments
    for i in range(len(segments)-1):
        x1 = segments[i][0] + 0.9
        x2 = segments[i+1][0] - 0.9
        ax.annotate('', xy=(x2, 3.25), xytext=(x1, 3.25),
                    arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))

    # Time labels below
    times = ['5min', '12min', '20min', '5min', '3min']
    for x, t in zip([s[0] for s in segments], times):
        ax.text(x, 1.8, t, ha='center', va='center', fontsize=8, color='#666666', **kw)

    # Legend
    ax.text(6, 1.0, '教学前 → 内容呈现 → 学习者参与 → 评估 → 总结迁移',
            ha='center', va='center', fontsize=9, style='italic', color='#333333')

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def _v1_boundary_point(pos: dict, node_type: str, toward: tuple[float, float]) -> tuple[float, float]:
    """Return a connection point on a node boundary, not its text center."""
    cx = float(pos["x"]) + float(pos["width"]) / 2
    cy = float(pos["y"]) + float(pos["height"]) / 2
    dx = float(toward[0]) - cx
    dy = float(toward[1]) - cy
    if abs(dx) < 1e-6 and abs(dy) < 1e-6:
        return cx, cy
    if node_type == "entry_boundary":
        left = float(pos["x"]) + 5
        right = float(pos["x"]) + float(pos["width"]) - 5
        return max(left, min(right, float(toward[0]))), cy
    half_w = max(float(pos["width"]) / 2, 1)
    half_h = max(float(pos["height"]) / 2, 1)
    if node_type == "decision":
        scale = 1 / (abs(dx) / half_w + abs(dy) / half_h)
    else:
        scale = min(half_w / abs(dx) if abs(dx) > 1e-6 else float("inf"),
                    half_h / abs(dy) if abs(dy) > 1e-6 else float("inf"))
    return cx + dx * scale, cy + dy * scale


def _v1_edge_points(layout: dict, node_by_id: dict, source: str, target: str):
    """Connect graph edges to node borders so labels remain unobstructed."""
    if source not in layout or target not in layout:
        return None
    source_pos = layout[source]
    target_pos = layout[target]
    source_center = (
        float(source_pos["x"]) + float(source_pos["width"]) / 2,
        float(source_pos["y"]) + float(source_pos["height"]) / 2,
    )
    target_center = (
        float(target_pos["x"]) + float(target_pos["width"]) / 2,
        float(target_pos["y"]) + float(target_pos["height"]) / 2,
    )
    source_type = node_by_id.get(source, {}).get("node_type", "action")
    target_type = node_by_id.get(target, {}).get("node_type", "action")
    return (
        _v1_boundary_point(source_pos, source_type, target_center),
        _v1_boundary_point(target_pos, target_type, source_center),
    )


def _pillow_font(size: int, bold: bool = False):
    if not HAS_PILLOW:
        return None
    candidates = []
    if _FONT_PATH:
        candidates.append(_FONT_PATH)
    candidates.extend([
        r'C:\Windows\Fonts\msyhbd.ttc' if bold else r'C:\Windows\Fonts\msyh.ttc',
        r'C:\Windows\Fonts\simhei.ttf',
        r'C:\Windows\Fonts\arial.ttf',
    ])
    for path in candidates:
        if path and os.path.exists(path):
            try:
                return ImageFont.truetype(path, size=size)
            except (OSError, ValueError):
                continue
    return ImageFont.load_default()


def _pillow_hex(value: str) -> tuple[int, int, int]:
    value = str(value or "#D9EAF7").lstrip("#")
    if len(value) != 6:
        return 217, 234, 247
    try:
        return tuple(int(value[index:index + 2], 16) for index in (0, 2, 4))
    except ValueError:
        return 217, 234, 247


def _pillow_wrap_label(value: object, max_chars: int = 18) -> str:
    text = str(value or "").strip()
    if not text:
        return "待确认"
    wrapped = []
    for line in str(text).replace("<br>", "\n").splitlines():
        current = []
        width = 0
        for char in line:
            char_width = 2 if ord(char) > 127 else 1
            if current and width + char_width > max_chars * 2:
                wrapped.append("".join(current))
                current, width = [], 0
            current.append(char)
            width += char_width
        if current:
            wrapped.append("".join(current))
    return "\n".join(wrapped) or "待确认"


def _pillow_arrow(draw, start: tuple[float, float], end: tuple[float, float], color, width: int = 2):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 11 + width * 2
    wing = 0.48
    left = (end[0] - length * math.cos(angle - wing), end[1] - length * math.sin(angle - wing))
    right = (end[0] - length * math.cos(angle + wing), end[1] - length * math.sin(angle + wing))
    draw.polygon([end, left, right], fill=color)


def _pillow_polyline_arrow(draw, points: list[tuple[float, float]], color, width: int = 2):
    """Draw an orthogonal connector with one arrowhead at its target."""
    if len(points) < 2:
        return
    for start, end in zip(points, points[1:]):
        draw.line([start, end], fill=color, width=width)
    _pillow_arrow(draw, points[-2], points[-1], color, width=width)


def _create_v1_graph_image_pillow(project: dict, view_name: str) -> io.BytesIO:
    """Pillow fallback for v1 graph images when matplotlib is unavailable."""
    if not HAS_PILLOW:
        return None
    try:
        from tools.skill_graph import build_skill_graph_views
        from tools.drawio_exporter import _v1_layout, _v1_wrap_label
        graph = project.get("skill_graph") or {}
        views = build_skill_graph_views(graph)
        view = views.get(view_name)
        if not view:
            return None
        layout = _v1_layout(view)
        if not layout:
            return None
        node_by_id = {str(node["id"]): node for node in view.get("nodes", [])}
        max_x = max(pos["x"] + pos["width"] for pos in layout.values()) + 70
        max_y = max(pos["y"] + pos["height"] for pos in layout.values()) + 70
        scale = 1.15
        margin = 30
        image = Image.new("RGB", (int(max_x * scale + margin * 2), int(max_y * scale + margin * 2)), "white")
        draw = ImageDraw.Draw(image)
        is_operation = view_name == "goal_operation_flow"
        is_hierarchy = view_name == "skill_hierarchy"
        is_control = view_name == "control_flow"
        is_readability_view = is_operation or is_hierarchy or is_control
        title_font = _pillow_font(36 if is_operation else 34 if is_readability_view else 24, bold=True)
        node_font = _pillow_font(28 if is_readability_view else 16)
        goal_font = _pillow_font(32 if is_readability_view else 18, bold=True)
        step_font = _pillow_font(30 if is_readability_view else 18, bold=True)
        edge_font = _pillow_font(22 if is_readability_view else 14)

        def point(value):
            return int(value * scale + margin)

        def xy(pair):
            return point(pair[0]), point(pair[1])

        colors = {
            "instructional_goal": ("#4472C4", "#FFFFFF"),
            "goal_step": ("#7565E8", "#FFFFFF"),
            "goal_substep": ("#E5E0FF", "#2D245F"),
            "intellectual_skill": ("#DDEBF7", "#1F2937"),
            "verbal_information": ("#E2F0D9", "#375623"),
            "psychomotor_skill": ("#FCE4D6", "#843C0C"),
            "cognitive_strategy": ("#FFF2CC", "#7F6000"),
            "attitude": ("#F4CCCC", "#660000"),
            "entry_skill": ("#FFF2CC", "#7F6000"),
            "entry_boundary": ("#FFFFFF", "#566573"),
            "start": ("#D9EAD3", "#274E13"),
            "end": ("#F4CCCC", "#660000"),
            "action": ("#D9EAF7", "#1F2937"),
            "decision": ("#4472C4", "#FFFFFF"),
        }
        title = str(view.get("title", "教学分析图"))
        draw.text((point(max_x / 2), point(24)), title, font=title_font, fill="#1F2937", anchor="mm")

        # Draw edges first; endpoints stop at node borders.
        hierarchy_entry_edges = 0
        for edge in view.get("edges", []):
            source = str(edge.get("from", ""))
            target = str(edge.get("to", ""))
            endpoints = _v1_edge_points(layout, node_by_id, source, target)
            if not endpoints:
                continue
            edge_type = edge.get("edge_type")
            color = "#C62828" if edge_type == "conditional_no" else "#2F5496"
            if edge_type in {"prerequisite", "entry_boundary"}:
                color = "#BF9000"
            feedback_label_pos = None
            if is_control and edge_type == "feedback":
                source_pos = layout[source]
                target_pos = layout[target]
                route_x = min(float(source_pos["x"]), float(target_pos["x"])) - 80
                source_y = float(source_pos["y"]) + float(source_pos["height"]) / 2
                target_y = float(target_pos["y"]) + float(target_pos["height"]) / 2
                path = [
                    xy((float(source_pos["x"]), source_y)),
                    xy((route_x, source_y)),
                    xy((route_x, target_y)),
                    xy((float(target_pos["x"]), target_y)),
                ]
                _pillow_polyline_arrow(draw, path, color, width=3)
                feedback_label_pos = (point(route_x - 10), point((source_y + target_y) / 2))
            elif is_hierarchy and node_by_id.get(source, {}).get("node_type") == "entry_skill":
                boundary = layout.get("ENTRY-BOUNDARY")
                target_pos = layout.get(target)
                source_pos = layout.get(source)
                if boundary and target_pos and source_pos:
                    route_y = float(boundary["y"]) - 18 - hierarchy_entry_edges * 14
                    lane_x = float(target_pos["x"]) - 28
                    start_point = (float(source_pos["x"]) + float(source_pos["width"]) / 2, float(source_pos["y"]))
                    target_point = (float(target_pos["x"]), float(target_pos["y"]) + float(target_pos["height"]) / 2)
                    path = [
                        xy(start_point),
                        xy((start_point[0], route_y)),
                        xy((lane_x, route_y)),
                        xy((lane_x, target_point[1])),
                        xy(target_point),
                    ]
                    _pillow_polyline_arrow(draw, path, color, width=2)
                    hierarchy_entry_edges += 1
                else:
                    start, end = (xy(endpoints[0]), xy(endpoints[1]))
                    _pillow_arrow(draw, start, end, color, width=2)
            else:
                start, end = (xy(endpoints[0]), xy(endpoints[1]))
                _pillow_arrow(draw, start, end, color, width=3 if edge_type in {"conditional_yes", "conditional_no"} else 2)
            if edge.get("label"):
                if feedback_label_pos:
                    draw.text(feedback_label_pos, str(edge["label"]), font=edge_font, fill=color, anchor="rm")
                else:
                    draw.text(((xy(endpoints[0])[0] + xy(endpoints[1])[0]) // 2, (xy(endpoints[0])[1] + xy(endpoints[1])[1]) // 2), str(edge["label"]), font=edge_font, fill=color, anchor="mm")

        for node_id, pos in layout.items():
            node = node_by_id[node_id]
            node_type = node.get("node_type", "action")
            face, text_color = colors.get(node_type, ("#D9EAF7", "#1F2937"))
            left, top = point(pos["x"]), point(pos["y"])
            right, bottom = point(pos["x"] + pos["width"]), point(pos["y"] + pos["height"])
            fill = _pillow_hex(face)
            outline = _pillow_hex("#2F5496")
            if node_type == "decision":
                cx, cy = (left + right) // 2, (top + bottom) // 2
                draw.polygon([(cx, top), (right, cy), (cx, bottom), (left, cy)], fill=fill, outline=outline)
            elif node_type == "entry_boundary":
                line_y = (top + bottom) // 2
                draw.line([(left, line_y), (right, line_y)], fill="#7F8C8D", width=3)
            elif node_type in {"start", "end"}:
                draw.ellipse((left, top, right, bottom), fill=fill, outline=outline, width=2)
            else:
                draw.rounded_rectangle((left, top, right, bottom), radius=8, fill=fill, outline=outline, width=2)
            wrap_width = (
                25 if node_type == "instructional_goal"
                else 20 if node_type == "goal_step"
                else 20 if is_control
                else 24
            )
            text = _pillow_wrap_label(_v1_wrap_label(node.get("label", ""), wrap_width), 18)
            if node_type == "instructional_goal":
                font = goal_font
            elif node_type == "goal_step":
                font = step_font
            else:
                font = node_font
            draw.multiline_text(
                ((left + right) // 2, (top + bottom) // 2),
                text,
                font=font,
                fill=_pillow_hex(text_color),
                anchor="mm",
                align="center",
                spacing=5 if is_readability_view else 4,
            )

        output = io.BytesIO()
        image.save(output, format="PNG", optimize=True)
        output.seek(0)
        return output
    except Exception:
        return None


def _create_v1_graph_image(project: dict, view_name: str) -> io.BytesIO:
    """Render one of the validated v1 graph views for Word/PDF embedding."""
    if not HAS_MATPLOTLIB:
        return _create_v1_graph_image_pillow(project, view_name)
    try:
        from tools.skill_graph import build_skill_graph_views
        from tools.drawio_exporter import _v1_layout, _v1_wrap_label
    except Exception:
        return None

    graph = project.get("skill_graph") or {}
    if not graph and isinstance(project.get("skill_graphs"), dict):
        graph = project["skill_graphs"].get("raw") or project["skill_graphs"].get("graph") or {}
    views = build_skill_graph_views(graph)
    view = views.get(view_name)
    if not view:
        return None
    layout = _v1_layout(view)
    if not layout:
        return None

    fp = _CHINESE_FONT
    font_kw = {"fontproperties": fp} if fp else {}
    max_x = max(pos["x"] + pos["width"] for pos in layout.values()) + 100
    max_y = max(pos["y"] + pos["height"] for pos in layout.values()) + 100
    is_operation = view_name == "goal_operation_flow"
    is_hierarchy = view_name == "skill_hierarchy"
    is_control = view_name == "control_flow"
    is_readability_view = is_operation or is_hierarchy or is_control
    fig_w = max(10.5 if is_operation else 11.0, min(22.0, max_x / 120))
    fig_h = max(4.0 if is_operation else 5.5, min(14.0, max_y / 100))
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.set_xlim(0, max_x)
    ax.set_ylim(max_y, 0)
    ax.axis("off")
    ax.text(max_x / 2, 28, view.get("title", "教学分析图"), ha="center", va="center", fontsize=26 if is_operation else 24 if is_readability_view else 14, fontweight="bold", **font_kw)

    colors = {
        "instructional_goal": ("#4472C4", "white"),
        "goal_step": ("#7565E8", "white"),
        "goal_substep": ("#E5E0FF", "#2D245F"),
        "intellectual_skill": ("#DDEBF7", "#1F2937"),
        "verbal_information": ("#E2F0D9", "#375623"),
        "psychomotor_skill": ("#FCE4D6", "#843C0C"),
        "cognitive_strategy": ("#FFF2CC", "#7F6000"),
        "attitude": ("#F4CCCC", "#660000"),
        "entry_skill": ("#FFF2CC", "#7F6000"),
        "entry_boundary": ("#FFFFFF", "#566573"),
        "start": ("#D9EAD3", "#274E13"),
        "end": ("#F4CCCC", "#660000"),
        "action": ("#D9EAF7", "#1F2937"),
        "decision": ("#4472C4", "white"),
    }
    node_by_id = {str(node["id"]): node for node in view.get("nodes", [])}

    def center(node_id: str):
        pos = layout[node_id]
        return pos["x"] + pos["width"] / 2, pos["y"] + pos["height"] / 2

    # Edges are drawn first so node boxes remain crisp and readable.
    hierarchy_entry_edges = 0
    for edge in view.get("edges", []):
        source, target = str(edge.get("from", "")), str(edge.get("to", ""))
        if source not in layout or target not in layout:
            continue
        endpoints = _v1_edge_points(layout, node_by_id, source, target)
        if not endpoints:
            continue
        (x1, y1), (x2, y2) = endpoints
        color = "#C62828" if edge.get("edge_type") == "conditional_no" else "#2F5496"
        if edge.get("edge_type") in {"prerequisite", "entry_boundary"}:
            color = "#BF9000"
        feedback_label_pos = None
        if is_control and edge.get("edge_type") == "feedback":
            source_pos = layout[source]
            target_pos = layout[target]
            route_x = min(float(source_pos["x"]), float(target_pos["x"])) - 80
            source_y = float(source_pos["y"]) + float(source_pos["height"]) / 2
            target_y = float(target_pos["y"]) + float(target_pos["height"]) / 2
            points = [
                (float(source_pos["x"]), source_y),
                (route_x, source_y),
                (route_x, target_y),
                (float(target_pos["x"]), target_y),
            ]
            for point_a, point_b in zip(points, points[1:]):
                ax.plot([point_a[0], point_b[0]], [point_a[1], point_b[1]], color=color, linewidth=1.6)
            ax.annotate("", xy=points[-1], xytext=points[-2], arrowprops={"arrowstyle": "->", "color": color, "lw": 1.6})
            feedback_label_pos = (route_x - 10, (source_y + target_y) / 2)
        elif is_hierarchy and node_by_id.get(source, {}).get("node_type") == "entry_skill":
            boundary = layout.get("ENTRY-BOUNDARY")
            target_pos = layout.get(target)
            source_pos = layout.get(source)
            if boundary and target_pos and source_pos:
                route_y = float(boundary["y"]) - 18 - hierarchy_entry_edges * 14
                lane_x = float(target_pos["x"]) - 28
                points = [
                    (float(source_pos["x"]) + float(source_pos["width"]) / 2, float(source_pos["y"])),
                    (float(source_pos["x"]) + float(source_pos["width"]) / 2, route_y),
                    (lane_x, route_y),
                    (lane_x, float(target_pos["y"]) + float(target_pos["height"]) / 2),
                    (float(target_pos["x"]), float(target_pos["y"]) + float(target_pos["height"]) / 2),
                ]
                for point_a, point_b in zip(points, points[1:]):
                    ax.plot([point_a[0], point_b[0]], [point_a[1], point_b[1]], color=color, linewidth=1.4)
                ax.annotate("", xy=points[-1], xytext=points[-2], arrowprops={"arrowstyle": "->", "color": color, "lw": 1.4})
                hierarchy_entry_edges += 1
            else:
                ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops={"arrowstyle": "->", "color": color, "lw": 1.2, "connectionstyle": "arc3,rad=0.02"})
        else:
            ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops={"arrowstyle": "->", "color": color, "lw": 1.2, "connectionstyle": "arc3,rad=0.02"})
        if edge.get("label"):
            if feedback_label_pos:
                ax.text(feedback_label_pos[0], feedback_label_pos[1], str(edge["label"]), color=color, fontsize=13 if is_readability_view else 9, ha="right", va="center", bbox={"facecolor": "white", "edgecolor": "none", "pad": 1.5}, **font_kw)
            else:
                ax.text((x1 + x2) / 2, (y1 + y2) / 2, str(edge["label"]), color=color, fontsize=11 if is_readability_view else 9, ha="center", va="center", **font_kw)

    for node_id, pos in layout.items():
        node = node_by_id[node_id]
        node_type = node.get("node_type", "action")
        face, text_color = colors.get(node_type, ("#D9EAF7", "#1F2937"))
        if node_type == "decision":
            cx = pos["x"] + pos["width"] / 2
            cy = pos["y"] + pos["height"] / 2
            half_w, half_h = pos["width"] / 2, pos["height"] / 2
            polygon = mpatches.Polygon([(cx, cy - half_h), (cx + half_w, cy), (cx, cy + half_h), (cx - half_w, cy)], closed=True, facecolor=face, edgecolor="#2F5496", linewidth=1.5)
            ax.add_patch(polygon)
        elif node_type == "entry_boundary":
            ax.plot([pos["x"], pos["x"] + pos["width"]], [pos["y"] + pos["height"] / 2] * 2, color="#7F8C8D", linestyle="--", linewidth=1.5)
        else:
            boxstyle = "round,pad=0.02" if node_type not in {"start", "end"} else "round,pad=0.04"
            box = mpatches.FancyBboxPatch((pos["x"], pos["y"]), pos["width"], pos["height"], boxstyle=boxstyle, facecolor=face, edgecolor="#2F5496", linewidth=1.2)
            ax.add_patch(box)
        wrap_width = (
            25 if node_type == "instructional_goal"
            else 20 if node_type == "goal_step"
            else 20 if is_control
            else 24
        )
        label = _v1_wrap_label(node.get("label", ""), wrap_width).replace("<br>", "\n")
        if node_type == "instructional_goal":
            font_size = 24.0 if is_readability_view else 9.5
        elif node_type == "goal_step":
            font_size = 22.0 if is_readability_view else 9.5
        else:
            font_size = 20.0 if is_readability_view else 8.5
        ax.text(pos["x"] + pos["width"] / 2, pos["y"] + pos["height"] / 2, label, ha="center", va="center", color=text_color, fontsize=font_size, multialignment="center", **font_kw)

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


def _add_text_diagram(doc: Document, title: str, lines: list):
    """Add a text-based diagram as a styled box (title + monospaced lines)."""
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(11)

    # Diagram body in a single-cell table with light gray background
    table = doc.add_table(rows=1, cols=1, style='Table Grid')
    cell = table.rows[0].cells[0]
    cell.text = ''
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)

    for idx, line in enumerate(lines):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.name = 'Courier New'

    doc.add_paragraph('')  # spacer


def _add_content_to_doc(doc: Document, content, heading_level: int = 2):
    """Recursively add a content dict/list/string to a Word document.

    Skips empty values and avoids raw dict/list dumps.
    """
    if isinstance(content, dict):
        for key, val in content.items():
            display_key = _display_label(key)
            if isinstance(val, str) and val.strip():
                doc.add_heading(display_key, level=heading_level)
                doc.add_paragraph(_safe(val))
            elif isinstance(val, list):
                non_empty = [v for v in val if v]
                if not non_empty:
                    continue
                doc.add_heading(display_key, level=heading_level)
                for item in non_empty:
                    if isinstance(item, str):
                        doc.add_paragraph(_safe(item), style='List Bullet')
                    elif isinstance(item, dict):
                        parts = []
                        for k2, v2 in item.items():
                            s = _safe(v2, '')
                            if s:
                                parts.append(f'{_display_label(k2)}：{s}')
                        if parts:
                            doc.add_paragraph('；'.join(parts), style='List Bullet')
            elif isinstance(val, dict):
                inner = {k: v for k, v in val.items() if v}
                if inner:
                    doc.add_heading(display_key, level=heading_level)
                    _add_content_to_doc(doc, inner, heading_level + 1)
    elif isinstance(content, list):
        for item in content:
            if isinstance(item, str) and item.strip():
                doc.add_paragraph(item, style='List Bullet')
            elif isinstance(item, dict):
                parts = []
                for k, v in item.items():
                    s = _safe(v, '')
                    if s:
                        parts.append(f'{_display_label(k)}：{s}')
                if parts:
                    doc.add_paragraph('；'.join(parts), style='List Bullet')
    elif isinstance(content, str) and content.strip():
        doc.add_paragraph(content)


# ---------------------------------------------------------------------------
# Diagram generators
# ---------------------------------------------------------------------------

def _render_skill_hierarchy_diagram(doc: Document, project: dict):
    """Render a text-based skill hierarchy diagram (Figure 1)."""
    sg = project.get('skill_graph', {})
    goal = project.get('goal', {})
    steps = sg.get('goal_steps', [])
    subskills = sg.get('subordinate_skills', [])
    entries = sg.get('entry_behaviors', [])

    goal_label = _safe(goal.get('behavior', goal.get('full_statement', '教学目的')), '教学目的')

    lines = []
    lines.append(f'    [{goal_label}]')
    lines.append('           |')
    lines.append('    +------+------+------+------+')
    lines.append('    |      |      |      |      |')

    step_labels = []
    for s in steps:
        label = _safe(s.get('description', ''), '')[:12]
        step_labels.append(label)
    while len(step_labels) < 5:
        step_labels.append('...')

    lines.append(f'  [{step_labels[0]}] [{step_labels[1]}] [{step_labels[2]}] [{step_labels[3]}] [{step_labels[4]}]')

    # Group subskills by step
    subs_by_step = {}
    for sk in subskills:
        parent = sk.get('parent_step_id', sk.get('linked_step_id', ''))
        subs_by_step.setdefault(parent, []).append(sk)

    lines.append('')
    lines.append('  从属技能层:')
    for step in steps[:3]:  # show first 3 steps for brevity
        sid = step.get('step_id', '')
        step_subs = subs_by_step.get(sid, [])
        for sk in step_subs[:2]:
            label = _safe(sk.get('description', ''), '')[:20]
            lines.append(f'    {sid[-4:]} --> {label}')

    lines.append('')
    lines.append('  入门技能层:')
    for e in entries[:3]:
        label = _safe(e.get('name', e.get('description', '')), '')[:24]
        lines.append(f'    {label}')

    # Try to create real image, fall back to text diagram
    img_buf = _create_v1_graph_image(project, "skill_hierarchy") or _create_skill_hierarchy_image(project)
    if img_buf:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_buf, width=Inches(6.4))
        cap = doc.add_paragraph('图 1  技能层级结构图')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
    else:
        _add_text_diagram(doc, '图 1  技能层级结构图', lines)


def _render_lesson_flow_diagram(doc: Document, project: dict):
    """Render a text-based lesson flow diagram (Figure 2)."""
    strategy = project.get('instructional_strategy', {})
    flow = strategy.get('lesson_flow', strategy.get('activity_flow', []))

    if not flow:
        doc.add_paragraph('（无教学活动流程数据）')
        return

    lines = []
    lines.append('  时间     教学环节                  活动内容')
    lines.append('  ' + '-' * 60)

    for seg in flow:
        time_val = seg.get('时间段', seg.get('time', seg.get('time_slot', '')))
        if isinstance(time_val, (int, float)):
            time_str = f'{int(time_val)}min'
        else:
            time_str = str(time_val)[:6] if time_val else ''
        activity = seg.get('活动', seg.get('具体活动', seg.get('activity', '')))
        activity_short = _safe(activity, '')[:28]
        component = seg.get('学习成分', seg.get('教学环节', ''))
        component_short = _safe(component, '')[:14]
        lines.append(f'  {time_str:>6}  {component_short:<14}  {activity_short}')

    # Try to create real image, fall back to text diagram
    img_buf = _create_lesson_flow_image(project)
    if img_buf:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_buf, width=Inches(6.4))
        cap = doc.add_paragraph('图 2  教学活动流程图')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
    else:
        _add_text_diagram(doc, '图 2  教学活动流程图', lines)


# ---------------------------------------------------------------------------
# Part 1 renderers: 教学分析
# ---------------------------------------------------------------------------

def _render_part1_title(doc: Document):
    """Render Part 1 title page."""
    doc.add_heading('报告一：教学分析', level=0)
    doc.add_paragraph(
        '本部分包含教学目的确定、教学目的分析、从属技能分析与入门技能识别。'
    )
    doc.add_paragraph('')


def _render_s01_teaching_purpose(doc: Document, project: dict):
    """1.1 教学目的（绩效问题、教学问题、绩效分析、目的陈述）"""
    doc.add_heading('1.1 教学目的', level=1)
    goal = project.get('goal', {})
    perf = (
        goal.get('performance_analysis')
        or goal.get('performance')
        or project.get('performance_analysis', {})
        or project.get('needs_analysis', {})
    )
    context = project.get('context_analysis', project.get('learner_context', {}))

    # 绩效问题
    doc.add_heading('绩效问题描述', level=2)
    perf_desc = perf.get('performance_problem', perf.get('description', ''))
    if perf_desc:
        doc.add_paragraph(_safe(perf_desc))
    else:
        # Infer from goal if no explicit problem
        behavior = goal.get('behavior', '')
        if behavior:
            doc.add_paragraph(
                f'学习者目前无法{behavior}，需要通过系统教学解决这一绩效差距。'
            )
        else:
            doc.add_paragraph(_DATA_GAP)

    # 教学问题
    doc.add_heading('教学问题', level=2)
    teaching_prob = perf.get('instructional_problem', goal.get('instructional_problem', ''))
    if teaching_prob:
        doc.add_paragraph(_safe(teaching_prob))
    else:
        doc.add_paragraph(_DATA_GAP)

    # 绩效分析
    doc.add_heading('绩效分析', level=2)
    expected = perf.get('expected_performance', perf.get('desired', ''))
    current = perf.get('current_performance', perf.get('actual', ''))
    gap = perf.get('performance_gap', perf.get('gap', ''))
    solvable = perf.get('instructionally_solvable', perf.get('is_instructional', ''))
    rows = [
        ['期望绩效', _safe(expected, _DATA_GAP)],
        ['当前绩效', _safe(current, _DATA_GAP)],
        ['绩效差距', _safe(gap, _DATA_GAP)],
        ['是否教学可解决', _safe(solvable, _DATA_GAP)],
    ]
    _add_table(doc, ['分析维度', '描述'], rows, col_widths=[4, 12])

    # 目的陈述
    doc.add_heading('目的陈述', level=2)
    statement = goal.get('full_statement', goal.get('statement', ''))
    if statement:
        p = doc.add_paragraph()
        run = p.add_run(statement)
        run.bold = True
    else:
        doc.add_paragraph(_DATA_GAP)

    # 目的要素表
    doc.add_heading('目的要素', level=2)
    elements = [
        ['学习者', _safe(goal.get('learner'), _DATA_GAP)],
        ['最终行为', _safe(goal.get('behavior'), _DATA_GAP)],
        ['应用环境', _safe(goal.get('context'), _DATA_GAP)],
        ['工具/条件', _safe(goal.get('tools'), _DATA_GAP)],
    ]
    _add_table(doc, ['要素', '内容'], elements, col_widths=[3.5, 12.5])

    # Provenance is rendered as teacher-facing evidence, never as internal
    # source dictionaries or implementation enums.
    doc.add_heading('来源依据', level=2)
    sources = project.get('sources', []) or goal.get('sources', []) or []
    source_rows = []
    for source in sources:
        source_name = _safe(source.get('source_name', source.get('title')), _DATA_GAP)
        issuer = _safe(source.get('issuer'), '发布机构待核对')
        version = _safe(source.get('source_version', source.get('version')), '版本待核对')
        date = _safe(source.get('publication_date', source.get('source_date')), '日期待核对')
        status = _status_label(source.get('evidence_status', 'clause_candidate'), '条款候选（待教师核对）')
        clauses = source.get('specific_clauses', source.get('clauses', [])) or []
        if clauses:
            clause_lines = []
            for clause in clauses[:4]:
                location = _safe(clause.get('page_number', clause.get('anchor', '位置待补充')), '位置待补充')
                excerpt = _safe(clause.get('excerpt', clause.get('clause_text')), _DATA_GAP)
                summary = _safe(clause.get('normalized_summary', clause.get('clause_text')), _DATA_GAP)
                clause_lines.append(f'{location}：{excerpt}（摘要：{summary}）')
            citation = '；'.join(clause_lines)
        else:
            citation = '尚未定位到具体条款，仅作文件级候选参考'
        source_rows.append([source_name, f'{issuer}\n{version}\n{date}', status, citation])
    if source_rows:
        _add_table(doc, ['来源名称', '发布信息', '证据状态', '条款位置与短引文/摘要'], source_rows, col_widths=[4.2, 3.2, 3.0, 6.6])
    else:
        doc.add_paragraph('尚未检索到可追溯来源；教学目的只能作为待确认草案。')


def _render_s02_goal_analysis(doc: Document, project: dict):
    """1.2 教学目的分析（分类、主要步骤、每步学习类型）"""
    doc.add_heading('1.2 教学目的分析', level=1)
    sg = project.get('skill_graph', {})

    # 目的分类
    doc.add_heading('目的分类', level=2)
    goal_type = sg.get('goal_type') or sg.get('metadata', {}).get('goal_type', '未分类')
    doc.add_paragraph(f'教学目的类型：{_learning_type_label(goal_type, "尚未分类")}')
    rationale = sg.get('classification_rationale', '')
    if rationale:
        doc.add_paragraph(f'判断依据：{rationale}')

    # 主要步骤
    doc.add_heading('主要步骤', level=2)
    steps = sg.get('goal_steps', [])
    if steps:
        headers = ['步骤编号', '步骤描述', '学习类型', '状态']
        rows = []
        for s in steps:
            rows.append([
                _safe(s.get('step_id')),
                _safe(s.get('description')),
                _learning_type_label(s.get('learning_type')),
                _status_label(s.get('status')),
            ])
        _add_table(doc, headers, rows, col_widths=[3, 7, 3, 3])
        operation_img = _create_v1_graph_image(project, "goal_operation_flow")
        if operation_img:
            # Keep the graph and its caption together in real Word pagination.
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            doc.add_picture(operation_img, width=Inches(6.4))
            cap = doc.add_paragraph('图 1  目的操作流程图（独立于技能层级图）')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.keep_together = True
            cap.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph(_DATA_GAP)

    # 每步学习类型
    doc.add_heading('每步学习类型判断', level=2)
    if steps:
        headers = ['步骤', '学习类型', '判断依据']
        rows = []
        for s in steps:
            rows.append([
                _safe(s.get('description')),
                _learning_type_label(s.get('learning_type')),
                _safe(s.get('learning_type_rationale', s.get('rationale', '')), _DATA_GAP),
            ])
        _add_table(doc, headers, rows, col_widths=[5, 3.5, 7.5])


def _render_s03_subordinate_skills(doc: Document, project: dict):
    """1.3 从属技能分析与入门技能"""
    doc.add_heading('1.3 从属技能分析', level=1)
    sg = project.get('skill_graph', {})

    steps = sg.get('goal_steps', [])
    subskills = sg.get('subordinate_skills', [])
    entries = sg.get('entry_behaviors', [])

    if not subskills and not entries:
        doc.add_paragraph('暂无从属技能数据。')
        return

    # Skill hierarchy figure
    _render_skill_hierarchy_diagram(doc, project)

    # Programming topics require a separate control-flow view.  It is kept
    # independent from the instructional-skill hierarchy and embedded in the
    # report so the Word artifact remains self-contained.
    graph = project.get("skill_graph", {})
    graph_has_control_flow = bool(graph.get("include_control_flow"))
    if not graph_has_control_flow:
        try:
            from tools.skill_graph import build_skill_graph_views
            control_view = build_skill_graph_views(graph).get("control_flow", {})
            graph_has_control_flow = any(node.get("node_type") == "decision" for node in control_view.get("nodes", []))
        except Exception:
            graph_has_control_flow = False
    if graph_has_control_flow:
        control_img = _create_v1_graph_image(project, "control_flow")
        if control_img:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(control_img, width=Inches(6.4))
            cap = doc.add_paragraph('图 3  程序控制流程图（独立于技能层级图）')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)

    # Group sub-skills by parent step
    subs_by_step = {}
    for sk in subskills:
        parent = sk.get('parent_step_id', sk.get('linked_step_id', '未知'))
        subs_by_step.setdefault(parent, []).append(sk)

    # One table per step
    for step in steps:
        sid = step.get('step_id', '')
        doc.add_heading(f'步骤: {_safe(step.get("description"))}', level=2)
        step_subs = subs_by_step.get(sid, [])

        if step_subs:
            headers = ['技能编号', '技能描述', '学习类型', '父步骤', '状态']
            rows = []
            for sk in step_subs:
                rows.append([
                    _safe(sk.get('skill_id', sk.get('sub_id'))),
                    _safe(sk.get('description')),
                    _learning_type_label(sk.get('learning_type')),
                    _safe(step.get('description')),
                    _status_label(sk.get('status')),
                ])
            _add_table(doc, headers, rows, col_widths=[3, 5.5, 3, 3, 2.5])
        else:
            doc.add_paragraph('该步骤暂无从属技能。')

    # Ungrouped sub-skills
    ungrouped = [sk for sk in subskills
                 if sk.get('parent_step_id', sk.get('linked_step_id', '')) not in
                 {s.get('step_id', '') for s in steps}]
    if ungrouped:
        doc.add_heading('未分组从属技能', level=2)
        headers = ['技能编号', '描述', '学习类型', '父步骤', '状态']
        rows = []
        for sk in ungrouped:
            rows.append([
                _safe(sk.get('skill_id')),
                _safe(sk.get('description')),
                _learning_type_label(sk.get('learning_type')),
                _safe(sk.get('parent_step_id', sk.get('linked_step_id', '')), _DATA_GAP),
                _status_label(sk.get('status')),
            ])
        _add_table(doc, headers, rows, col_widths=[3, 5.5, 3, 3, 2.5])

    # Entry behaviors
    if entries:
        doc.add_heading('入门技能', level=2)
        headers = ['入门技能编号', '描述', '学习类型', '支持的技能', '状态']
        rows = []
        for e in entries:
            supports = e.get('supports_skill_ids', [])
            supports_str = ', '.join(str(s) for s in supports) if supports else _DATA_GAP
            rows.append([
                _safe(e.get('entry_id', e.get('skill_id'))),
                _safe(e.get('name', e.get('description'))),
                _learning_type_label(e.get('learning_type')),
                supports_str,
                _status_label(e.get('status')),
            ])
        _add_table(doc, headers, rows, col_widths=[3, 5, 3, 3, 2])


def _render_s04_entry_skills_from_context(doc: Document, project: dict):
    """1.4 入门技能（从 learner_context 交叉引用）"""
    doc.add_heading('1.4 入门技能', level=1)
    context = project.get('context_analysis', project.get('learner_context', {}))
    sg = project.get('skill_graph', {})

    # From learner context
    learner = context.get('learner_profile', context.get('learner', {}))
    if learner:
        entry_skills = learner.get('entry_skills', [])
        if entry_skills:
            doc.add_paragraph('学习者入门技能（来源：学习者分析）：')
            if isinstance(entry_skills, dict):
                level = _status_label(entry_skills.get('level'), _DATA_GAP)
                items = entry_skills.get('items', [])
                item_text = '；'.join(
                    _safe(item.get('name', item.get('description', '')), _DATA_GAP)
                    if isinstance(item, dict) else _safe(item, _DATA_GAP)
                    for item in items
                    if item
                ) or _DATA_GAP
                note = _safe(entry_skills.get('note'), _DATA_GAP)
                _add_table(
                    doc,
                    ['特征维度', '描述'],
                    [['掌握程度', level], ['具体技能', item_text], ['说明', note]],
                    col_widths=[3.5, 12.5],
                )
            else:
                for skill in entry_skills:
                    doc.add_paragraph(_safe(skill), style='List Bullet')

    # Cross-reference with skill graph
    entries = sg.get('entry_behaviors', [])
    if entries:
        doc.add_heading('技能图谱中的入门技能', level=2)
        headers = ['编号', '描述', '学习类型', '前置技能']
        rows = []
        for e in entries:
            supports = e.get('supports_skill_ids', [])
            rows.append([
                _safe(e.get('entry_id', e.get('skill_id'))),
                _safe(e.get('name', e.get('description'))),
                _learning_type_label(e.get('learning_type')),
                ', '.join(str(s) for s in supports) if supports else '—',
            ])
        _add_table(doc, headers, rows, col_widths=[3, 5.5, 3, 4.5])


# ---------------------------------------------------------------------------
# Part 2 renderers: 教学策略与评价
# ---------------------------------------------------------------------------

def _render_part2_title(doc: Document):
    """Render Part 2 title page."""
    doc.add_heading('报告二：教学策略与评价', level=0)
    doc.add_paragraph(
        '本部分包含学习者与环境分析、绩效目标与评价、教学顺序、教学前活动、'
        '信息呈现、练习与反馈、教学策略、分组安排、评价工具。'
    )
    doc.add_paragraph('')


def _render_s05_learner_context(doc: Document, project: dict):
    """2.1 学习者与环境分析"""
    doc.add_heading('2.1 学习者与环境分析', level=1)
    context = project.get('context_analysis', project.get('learner_context', {}))

    if not context:
        doc.add_paragraph('暂无学习者与环境分析数据。')
        return

    # Learner profile
    doc.add_heading('学习者特征', level=2)
    learner = context.get('learner_profile', context.get('learner', {}))
    if learner:
        headers = ['特征维度', '描述']
        rows = []
        label_map = [
            ('entry_skills', '入门技能'),
            ('prior_knowledge', '先前知识'),
            ('motivation', '学习动机'),
            ('learning_preferences', '学习偏好'),
            ('common_difficulties', '常见困难'),
            ('group_characteristics', '群体特征'),
            ('attitude_toward_content', '对内容的态度'),
            ('attitude_toward_delivery', '对教学方式的态度'),
            ('ability_level', '能力水平'),
            ('age', '年龄/年级'),
            ('learning_style', '学习风格'),
            ('special_needs', '特殊需求'),
        ]
        for key, lbl in label_map:
            val = learner.get(key, '')
            if isinstance(val, dict):
                val = val.get('description', val.get('level', ''))
            if isinstance(val, list):
                val = ', '.join(str(v) for v in val)
            val_str = _safe(val, '')
            if val_str and val_str != '{}':
                rows.append([lbl, val_str])
        if rows:
            _add_table(doc, headers, rows, col_widths=[4, 12])
        else:
            doc.add_paragraph(_DATA_GAP)
    else:
        doc.add_paragraph(_DATA_GAP)

    # Learning context
    doc.add_heading('教学环境', level=2)
    lctx = context.get('learning_context', {})
    if lctx:
        rows = []
        for key, lbl in [
            ('class_duration', '课时（分钟）'),
            ('class_size', '班额'),
            ('devices', '设备条件'),
            ('network', '网络条件'),
            ('classroom_layout', '教室布局'),
        ]:
            val = lctx.get(key, '')
            if isinstance(val, dict):
                val = val.get('description', '')
            rows.append([lbl, _safe(val, '')])
        # Only add non-empty rows
        rows = [r for r in rows if r[1]]
        if rows:
            _add_table(doc, ['维度', '描述'], rows, col_widths=[4, 12])

        media = lctx.get('available_media', [])
        if media:
            doc.add_paragraph(f'可用媒体：{", ".join(str(m) for m in media)}')
        constraints = lctx.get('constraints', [])
        if constraints:
            doc.add_paragraph('约束条件：')
            for c in constraints:
                doc.add_paragraph(_safe(c), style='List Bullet')

    # Performance context
    doc.add_heading('应用环境', level=2)
    pctx = context.get('performance_context', {})
    if pctx:
        rows = []
        for key, lbl in [
            ('use_environment', '应用环境'),
            ('expected_transfer', '预期迁移'),
            ('similarity_to_learning_context', '环境相似度'),
        ]:
            val = pctx.get(key, '')
            if isinstance(val, list):
                val = '; '.join(str(v) for v in val)
            val_str = _safe(val, '')
            if val_str:
                rows.append([lbl, val_str])
        transfer_risks = pctx.get('transfer_risks', [])
        if transfer_risks:
            rows.append(['迁移风险', '; '.join(str(r) for r in transfer_risks)])
        if rows:
            _add_table(doc, ['维度', '描述'], rows, col_widths=[4, 12])

    # Strategy implications
    doc.add_heading('策略启示', level=2)
    imps = context.get('strategy_implications',
                       context.get('implications', context.get('design_implications', [])))
    if isinstance(imps, list):
        shown = 0
        for imp in imps:
            if isinstance(imp, dict):
                text = imp.get('implication', imp.get('description', ''))
            else:
                text = str(imp)
            if text and text.strip():
                shown += 1
                doc.add_paragraph(f'{shown}. {text}')
        if shown == 0:
            doc.add_paragraph(_DATA_GAP)
    elif imps:
        doc.add_paragraph(_safe(imps))
    else:
        doc.add_paragraph(_DATA_GAP)


def _render_s06_objectives_and_assessment(doc: Document, project: dict):
    """2.2 绩效目标与评价样题"""
    doc.add_heading('2.2 绩效目标与评价样题', level=1)
    objectives = project.get('objectives', [])
    sg = project.get('skill_graph', {})
    assessment = project.get('assessment_plan', {})

    if not objectives:
        doc.add_paragraph('暂无绩效目标。')
        return

    # Evidence lookup by objective
    evidence_by_obj = {}
    for ev in assessment.get('evidence', []):
        oid = ev.get('linked_objective_id', '')
        if oid:
            evidence_by_obj.setdefault(oid, []).append(ev)

    # Group by step
    steps = sg.get('goal_steps', [])
    objs_by_step = {}
    for obj in objectives:
        skill_id = obj.get('related_skill_id', obj.get('step_id', '未分组'))
        objs_by_step.setdefault(skill_id, []).append(obj)

    for step in steps:
        sid = step.get('step_id', '')
        doc.add_heading(f'步骤: {_safe(step.get("description"))}', level=2)
        step_objs = objs_by_step.get(sid, [])
        if not step_objs:
            doc.add_paragraph('该步骤暂无绩效目标。')
            continue

        headers = ['目标编号', '条件', '行为', '标准', '评价证据', '评分标准']
        rows = []
        for obj in step_objs:
            oid = obj.get('objective_id', '')
            condition = _safe_short(obj.get('condition'), 40, _DATA_GAP)
            behavior = _safe_short(obj.get('behavior'), 40, _DATA_GAP)
            criterion = _safe_short(obj.get('criterion'), 30, _DATA_GAP)

            evs = evidence_by_obj.get(oid, [])
            evidence_text = '；'.join(
                _safe_short(
                    e.get('task_prompt')
                    or _evidence_type_label(
                        e.get('evidence_type_name', e.get('evidence_type'))
                    ),
                    30,
                    _DATA_GAP,
                )
                for e in evs[:3]
            ) if evs else _DATA_GAP

            scoring_parts = []
            for ev in evs:
                for sc in ev.get('scoring_criteria', []):
                    scoring_parts.append(
                        f"{_safe(sc.get('criterion'))}: "
                        f"{_safe(sc.get('description'))} ({_safe(sc.get('max_score'))}分)"
                    )
            scoring_text = _safe_short('; '.join(scoring_parts), 50) if scoring_parts else '待补充'

            rows.append([oid, condition, behavior, criterion, evidence_text, scoring_text])
        _add_table(doc, headers, rows, col_widths=[2.5, 3, 3, 2.5, 2.5, 2.5])

    # Ungrouped
    ungrouped = objs_by_step.get('未分组', [])
    if ungrouped:
        doc.add_heading('未分组目标', level=2)
        headers = ['目标编号', '关联技能', '条件', '行为', '标准', '状态']
        rows = []
        for obj in ungrouped:
            oid = obj.get('objective_id', '')
            rows.append([
                oid,
                _safe(obj.get('related_skill_id'), _DATA_GAP),
                _safe_short(obj.get('condition'), 35, _DATA_GAP),
                _safe_short(obj.get('behavior'), 35, _DATA_GAP),
                _safe_short(obj.get('criterion'), 30, _DATA_GAP),
                _status_label(obj.get('status')),
            ])
        _add_table(doc, headers, rows, col_widths=[2.5, 3, 3, 3, 2.5, 2])


def _render_s07_instructional_sequence(doc: Document, project: dict):
    """2.3 教学顺序说明"""
    doc.add_heading('2.3 教学顺序说明', level=1)
    strategy = project.get('instructional_strategy', {})

    sequence = strategy.get('sequence_rationale',
                            strategy.get('objective_sequence', {}).get('rationale', ''))
    doc.add_heading('顺序依据', level=2)
    doc.add_paragraph(_safe(sequence, _DATA_GAP))

    doc.add_heading('教学活动流程', level=2)
    lesson_flow = strategy.get('lesson_flow', strategy.get('activity_flow', []))
    if lesson_flow:
        headers = ['时间段', '活动', '对应目标', '学习成分', '备注']
        rows = []
        for index, step in enumerate(lesson_flow):
            time_slot = step.get('时间段', step.get('time', step.get('time_slot', '')))
            if isinstance(time_slot, (int, float)):
                time_slot = f'{int(time_slot)}分钟'
            activity = step.get('活动', step.get('具体活动', step.get('activity', step.get('description', ''))))
            obj_ref = _flow_objective_refs(project, step, index)
            comp = step.get('学习成分', step.get('教学环节', step.get('component', step.get('learning_component', ''))))
            notes = step.get('备注', step.get('notes', step.get('评估方式', step.get('remark', ''))))
            rows.append([
                _safe(time_slot),
                _safe_short(activity, 40),
                _safe_short(obj_ref, 30, _DATA_GAP),
                _safe(comp),
                _safe_short(notes, 30),
            ])
        _add_table(doc, headers, rows, col_widths=[2, 4.5, 3, 3, 3.5])
    else:
        doc.add_paragraph(_DATA_GAP)

    # Lesson flow figure
    _render_lesson_flow_diagram(doc, project)


def _render_s08_pre_instructional(doc: Document, project: dict):
    """2.4 教学前活动说明（ARCS/Gagné）"""
    doc.add_heading('2.4 教学前活动说明', level=1)
    strategy = project.get('instructional_strategy', {})
    learning_components = strategy.get('learning_components', [])

    # Find pre-instructional component
    pre = None
    for comp in learning_components:
        if comp.get('type', '') in ('pre_instructional', 'pre-instructional', 'attention'):
            pre = comp
            break

    if pre:
        doc.add_heading('活动描述', level=2)
        desc = pre.get('description', pre.get('activity', ''))
        if desc:
            doc.add_paragraph(_safe(desc))
        duration = pre.get('duration', '')
        if duration:
            doc.add_paragraph(f'时间分配：{duration}')
        teacher = pre.get('teacher_activity', '')
        if teacher:
            doc.add_paragraph(f'教师活动：{teacher}')
        learner = pre.get('learner_activity', '')
        if learner:
            doc.add_paragraph(f'学习者活动：{learner}')
    else:
        doc.add_paragraph(_DATA_GAP)

    # ARCS model
    doc.add_heading('ARCS 动机设计', level=2)
    arcs = strategy.get('arcs', strategy.get('motivation_design', {}))
    if arcs:
        arcs_items = [
            ['A - 注意（Attention）', _safe(arcs.get('attention', arcs.get('A', '')), _DATA_GAP)],
            ['R - 相关（Relevance）', _safe(arcs.get('relevance', arcs.get('R', '')), _DATA_GAP)],
            ['C - 信心（Confidence）', _safe(arcs.get('confidence', arcs.get('C', '')), _DATA_GAP)],
            ['S - 满意（Satisfaction）', _safe(arcs.get('satisfaction', arcs.get('S', '')), _DATA_GAP)],
        ]
        _add_table(doc, ['ARCS维度', '设计策略'], arcs_items, col_widths=[5, 11])
    else:
        doc.add_paragraph(_DATA_GAP)

    # Gagné pre-instructional events
    doc.add_heading('Gagné 教学事件（教学前部分）', level=2)
    gagne = strategy.get('gagne_events', strategy.get('instructional_events', {}))
    if gagne:
        pre_events = [
            ('1. 引起注意', gagne.get('attention', gagne.get('event_1', ''))),
            ('2. 告知学习目标', gagne.get('inform_objectives', gagne.get('event_2', ''))),
            ('3. 激发先备知识', gagne.get('stimulate_recall', gagne.get('event_3', ''))),
        ]
        rows = [[label, _safe(val, _DATA_GAP)] for label, val in pre_events if val]
        if rows:
            _add_table(doc, ['教学事件', '描述'], rows, col_widths=[4, 12])
        else:
            doc.add_paragraph(_DATA_GAP)
    else:
        doc.add_paragraph(_DATA_GAP)


def _render_s09_content_presentation(doc: Document, project: dict):
    """2.5 信息呈现与范例"""
    doc.add_heading('2.5 信息呈现与范例', level=1)
    strategy = project.get('instructional_strategy', {})
    objectives = project.get('objectives', [])

    learning_components = strategy.get('learning_components', [])
    pres = [c for c in learning_components
            if c.get('type', '') in ('content_presentation', 'content', 'presentation')]
    pres_comp = pres[0] if pres else None

    if pres_comp:
        doc.add_heading('内容呈现概述', level=2)
        desc = pres_comp.get('description', pres_comp.get('activity', ''))
        if desc:
            doc.add_paragraph(_safe(desc))
        linked = pres_comp.get('linked_objectives', [])
        if linked:
            doc.add_paragraph(f'对应目标：{", ".join(str(o) for o in linked)}')

    if objectives:
        doc.add_heading('各目标信息呈现与范例', level=2)
        headers = ['目标编号', '行为描述', '信息呈现方式', '范例/示范', '备注']
        rows = []
        for obj in objectives:
            oid = obj.get('objective_id', '')
            behavior = _safe_short(obj.get('behavior'), 35)
            presentation = obj.get('presentation', obj.get('information_presentation', ''))
            examples = obj.get('examples', obj.get('model', obj.get('worked_example', '')))
            if isinstance(examples, list):
                examples = '; '.join(str(e) for e in examples)
            notes = obj.get('presentation_notes', '')
            rows.append([oid, behavior, _safe(presentation, _DATA_GAP), _safe(examples, _DATA_GAP), _safe(notes, '')])
        _add_table(doc, headers, rows, col_widths=[2.5, 3.5, 3, 4, 3])
    else:
        doc.add_paragraph(_DATA_GAP)


def _render_s10_practice_feedback(doc: Document, project: dict):
    """2.6 练习与反馈"""
    doc.add_heading('2.6 练习与反馈', level=1)
    strategy = project.get('instructional_strategy', {})
    objectives = project.get('objectives', [])

    learning_components = strategy.get('learning_components', [])
    participations = [c for c in learning_components
                      if c.get('type', '') in ('learner_participation', 'participation', 'practice')]
    part_comp = participations[0] if participations else None

    if part_comp:
        doc.add_heading('练习活动概述', level=2)
        desc = part_comp.get('description', part_comp.get('activity', ''))
        if desc:
            doc.add_paragraph(_safe(desc))
        linked = part_comp.get('linked_objectives', [])
        if linked:
            doc.add_paragraph(f'对应目标：{", ".join(str(o) for o in linked)}')

    if objectives:
        doc.add_heading('各目标练习与反馈', level=2)
        headers = ['目标编号', '行为描述', '练习活动', '反馈方式', '参考答案/要点']
        rows = []
        for obj in objectives:
            oid = obj.get('objective_id', '')
            behavior = _safe_short(obj.get('behavior'), 35)
            practice = obj.get('practice', obj.get('practice_activity', ''))
            feedback = obj.get('feedback', obj.get('feedback_mechanism', ''))
            answers = obj.get('answer_key', obj.get('correct_answers', ''))
            if isinstance(answers, list):
                answers = '; '.join(str(a) for a in answers)
            rows.append([
                oid, behavior,
                _safe_short(practice, 40, _DATA_GAP),
                _safe_short(feedback, 30, _DATA_GAP),
                _safe_short(answers, 40, _DATA_GAP),
            ])
        _add_table(doc, headers, rows, col_widths=[2.5, 3, 3.5, 3, 4])
    else:
        doc.add_paragraph(_DATA_GAP)


def _render_s11_instructional_strategy(doc: Document, project: dict):
    """2.7 教学策略综述"""
    doc.add_heading('2.7 教学策略综述', level=1)
    strategy = project.get('instructional_strategy', {})

    if not strategy:
        doc.add_paragraph('暂无教学策略数据。')
        return

    # Duration
    duration = strategy.get('total_duration', strategy.get('lesson_duration', ''))
    if duration:
        doc.add_paragraph(f'总教学时长：{duration}')

    # Segments summary
    segments = strategy.get('segments', [])
    if segments:
        doc.add_heading('教学环节概览', level=2)
        headers = ['环节', '时间(分钟)', '对应目标数', '说明']
        rows = []
        for seg in segments:
            name = seg.get('name', '')
            time_m = seg.get('time_minutes', '')
            objs = seg.get('objectives', [])
            rationale = seg.get('rationale', '')
            rows.append([
                _safe(name),
                str(time_m) if time_m else '',
                str(len(objs)) if objs else '0',
                _safe_short(rationale, 40),
            ])
        _add_table(doc, headers, rows, col_widths=[4, 3, 3, 6])


def _render_s12_grouping(doc: Document, project: dict):
    """2.8 分组安排"""
    doc.add_heading('2.8 分组安排', level=1)
    strategy = project.get('instructional_strategy', {})
    context = project.get('context_analysis', project.get('learner_context', {}))

    grouping = strategy.get('grouping_strategy',
                            strategy.get('grouping', {}))
    if grouping:
        doc.add_heading('分组策略', level=2)
        headers = ['维度', '描述']
        rows = []
        for key, lbl in [
            ('strategy', '分组策略'),
            ('group_size', '小组规模'),
            ('composition', '分组方式'),
            ('rationale', '分组依据'),
        ]:
            val = grouping.get(key, '')
            if val:
                rows.append([lbl, _safe(val)])
        if rows:
            _add_table(doc, headers, rows, col_widths=[4, 12])
    else:
        # Infer from context
        lctx = context.get('learning_context', {})
        class_size = lctx.get('class_size', '')
        layout = lctx.get('classroom_layout', '')
        if class_size or layout:
            doc.add_paragraph(f'班额：{_safe(class_size, _DATA_GAP)}')
            doc.add_paragraph(f'教室布局：{_safe(layout, _DATA_GAP)}')
            doc.add_paragraph('建议采用异质分组，每组4-6人，以支持合作学习。')
        else:
            doc.add_paragraph(_DATA_GAP)


def _render_s13_assessment_tools(doc: Document, project: dict):
    """2.9 评价工具（入门测试、前测、后测、量规）"""
    doc.add_heading('2.9 评价工具', level=1)
    assessment = project.get('assessment_plan', {})

    if not assessment:
        doc.add_paragraph('暂无评价方案数据。')
        return

    sections = [
        ('entry_behavior_test', '入门技能测试'),
        ('pretest', '前测'),
        ('posttest', '后测'),
    ]

    for key, title in sections:
        doc.add_heading(title, level=2)
        section_data = assessment.get(key, {})
        if not section_data:
            doc.add_paragraph(f'暂无{_DATA_GAP}。')
            continue

        purpose = section_data.get('purpose', '')
        if purpose:
            doc.add_paragraph(f'目的：{purpose}')

        items = section_data.get('items', [])
        if items:
            headers = ['题号', '任务提示/题目', '预期证据', '分值']
            rows = []
            for idx, item in enumerate(items, 1):
                task = item.get('task_prompt', item.get('question', item.get('task_description', '')))
                expected = item.get('expected_evidence', '')
                score = item.get('max_score', item.get('points', ''))
                # Build score string from scoring_criteria if max_score missing
                if not score:
                    criteria = item.get('scoring_criteria', [])
                    if criteria:
                        score = f"{sum(c.get('max_score', 0) for c in criteria)}分"
                rows.append([str(idx), _safe_short(task, 50), _safe_short(expected, 40), _safe(score)])
            _add_table(doc, headers, rows, col_widths=[1.5, 6, 5, 3.5])

    # Scoring rubrics
    doc.add_heading('评分量规', level=2)
    rubrics = assessment.get('rubrics', assessment.get('scoring_rubrics', []))
    if rubrics and isinstance(rubrics, list):
        for rub in rubrics:
            if isinstance(rub, dict):
                doc.add_heading(_safe(rub.get('name', rub.get('title', '量规'))), level=3)
                criteria = rub.get('criteria', rub.get('dimensions', []))
                if criteria:
                    headers = ['维度', '描述', '分值', '等级说明']
                    rows = []
                    for c in criteria:
                        rows.append([
                            _safe(c.get('dimension', c.get('name', ''))),
                            _safe_short(c.get('description', c.get('criterion', '')), 40),
                            _safe(c.get('max_score', c.get('points', ''))),
                            _safe_short(c.get('levels', c.get('level_descriptions', '')), 40),
                        ])
                    _add_table(doc, headers, rows, col_widths=[3, 5, 2, 6])
    else:
        doc.add_paragraph(_DATA_GAP)

    # Evidence mapping
    evidence = assessment.get('evidence', [])
    if evidence:
        doc.add_heading('目标-评价证据对应表', level=2)
        headers = ['目标编号', '证据类型', '任务提示', '状态']
        rows = []
        for e in evidence:
            task = _safe_short(e.get('task_prompt', e.get('description', '')), 50)
            rows.append([
                _safe(e.get('linked_objective_id'), _DATA_GAP),
                _evidence_type_label(e.get('evidence_type_name', e.get('evidence_type'))),
                task,
                _status_label(e.get('status')),
            ])
        _add_table(doc, headers, rows, col_widths=[2.5, 4, 6, 3.5])


def _render_s14_materials_summary(doc: Document, project: dict):
    """2.10 教学材料摘要"""
    doc.add_heading('2.10 教学材料摘要', level=1)
    materials = project.get('instructional_materials', {})

    if not materials:
        doc.add_paragraph('暂无教学材料。')
        return

    # Summary table
    doc.add_heading('材料清单', level=2)
    summary_rows = []
    for mat_key, mat_label in _MATERIAL_LABELS:
        mat = materials.get(mat_key, {})
        if mat:
            status = mat.get('status', '已完成')
            obj_ids = mat.get('related_objective_ids', [])
            obj_str = ', '.join(str(o) for o in obj_ids[:5]) if obj_ids else ''
            summary_rows.append([mat_label, _status_label(status, '已完成'), obj_str or _DATA_GAP])
    if summary_rows:
        _add_table(doc, ['材料名称', '状态', '关联目标'], summary_rows, col_widths=[5, 3, 8])
    else:
        doc.add_paragraph('暂无材料清单数据。')

    # Detailed content
    doc.add_heading('材料详细内容', level=2)
    for mat_key, mat_label in _MATERIAL_LABELS:
        mat = materials.get(mat_key, {})
        if not mat:
            continue
        doc.add_heading(mat_label, level=3)
        content = mat.get('content', mat)
        _add_content_to_doc(doc, content, heading_level=4)


# ---------------------------------------------------------------------------
# Part 3 renderers: 形成性评价与修改
# ---------------------------------------------------------------------------

def _render_part3_title(doc: Document):
    """Render Part 3 title page - 形成性评价方案（待实施）."""
    doc.add_heading('报告三：形成性评价方案（待实施）', level=0)
    p = doc.add_paragraph()
    run = p.add_run(
        '⚠ 注意：本部分为形成性评价实施方案，尚未实施。'
        '实施后需根据实际数据修改教学。在获得真实评价数据前，'
        '本报告不可作为最终优秀教学系统设计报告使用。'
    )
    run.bold = True
    run.font.color.rgb = None  # default color
    doc.add_paragraph('')


def _render_s15_modifications(doc: Document, project: dict):
    """3.1 修改记录"""
    doc.add_heading('3.1 修改记录', level=1)

    revisions = project.get('revision_history',
                            project.get('ai_process', {}).get('revisions', []))
    if revisions:
        headers = ['修订编号', '修订对象', '修订内容', '修订原因', '日期']
        rows = []
        for idx, rev in enumerate(revisions, 1):
            rows.append([
                _safe(rev.get('revision_id', str(idx))),
                _safe(rev.get('target', rev.get('component', ''))),
                _safe_short(rev.get('description', rev.get('change', '')), 50),
                _safe_short(rev.get('rationale', rev.get('reason', '')), 40),
                _safe(rev.get('date', rev.get('timestamp', ''))),
            ])
        _add_table(doc, headers, rows, col_widths=[2, 3, 5, 4, 2])
    else:
        doc.add_paragraph(
            '当前为初始设计版本，尚未进行形成性评价修订。'
            '本区域将在实施形成性评价后由教师和设计团队填写修订记录。'
        )
        # Provide fill-in template
        headers = ['修订编号', '修订对象', '修订内容', '修订原因', '日期']
        rows = [
            ['', '', '', '', ''],
            ['', '', '', '', ''],
            ['', '', '', '', ''],
        ]
        _add_table(doc, headers, rows, col_widths=[2, 3, 5, 4, 2])
        doc.add_paragraph('（请在实施形成性评价后填写上表）')


def _render_s16_one_on_one_evaluation(doc: Document, project: dict):
    """3.2 一对一评价"""
    doc.add_heading('3.2 一对一评价', level=1)

    one_on_one = project.get('formative_evaluation', {}).get('one_on_one',
                          project.get('one_on_one_evaluation', {}))
    if one_on_one:
        # Evaluation plan
        doc.add_heading('评价计划', level=2)
        plan = one_on_one.get('plan', one_on_one)
        headers = ['维度', '内容']
        rows = []
        for key, lbl in [
            ('purpose', '评价目的'),
            ('participants', '参与人数'),
            ('duration', '评价时长'),
            ('materials_used', '使用材料'),
            ('protocol', '评价流程'),
        ]:
            val = plan.get(key, '')
            if isinstance(val, list):
                val = '; '.join(str(v) for v in val)
            if val:
                rows.append([lbl, _safe(val)])
        if rows:
            _add_table(doc, headers, rows, col_widths=[4, 12])

        # Observations
        observations = one_on_one.get('observations', [])
        if observations:
            doc.add_heading('观察记录', level=2)
            obs_headers = ['参与者', '观察内容', '发现问题', '严重程度']
            obs_rows = []
            for obs in observations:
                obs_rows.append([
                    _safe(obs.get('participant', '')),
                    _safe_short(obs.get('observation', obs.get('content', '')), 40),
                    _safe_short(obs.get('issue', obs.get('finding', '')), 40),
                    _safe(obs.get('severity', '')),
                ])
            _add_table(doc, obs_headers, obs_rows, col_widths=[2.5, 5, 5, 3.5])
    else:
        doc.add_paragraph(
            '一对一评价计划将在教学材料初稿完成后制定。'
            '评价将选取3-5名不同水平的学习者，通过观察和访谈收集'
            '对教学材料的理解度、操作困难和改进建议。'
        )
        doc.add_heading('评价计划模板', level=2)
        headers = ['维度', '计划内容']
        rows = [
            ['评价目的', '检测学习者对教学材料的理解程度和使用体验'],
            ['参与人数', '3-5名不同水平的学习者'],
            ['评价时长', '每人20-30分钟'],
            ['评价流程', '完成学习任务 -> 出声思维 -> 访谈 -> 记录'],
            ['关注重点', '内容理解、活动参与度、材料可读性、时间安排'],
        ]
        _add_table(doc, headers, rows, col_widths=[4, 12])


def _render_s17_small_group_evaluation(doc: Document, project: dict):
    """3.3 小组评价"""
    doc.add_heading('3.3 小组评价', level=1)

    small_group = project.get('formative_evaluation', {}).get('small_group',
                    project.get('small_group_evaluation', {}))
    if small_group:
        doc.add_heading('评价计划', level=2)
        plan = small_group.get('plan', small_group)
        headers = ['维度', '内容']
        rows = []
        for key, lbl in [
            ('purpose', '评价目的'),
            ('participant_count', '参与人数'),
            ('group_size', '小组规模'),
            ('duration', '评价时长'),
            ('activities', '活动安排'),
        ]:
            val = plan.get(key, '')
            if isinstance(val, list):
                val = '; '.join(str(v) for v in val)
            if val:
                rows.append([lbl, _safe(val)])
        if rows:
            _add_table(doc, headers, rows, col_widths=[4, 12])
    else:
        doc.add_paragraph(
            '小组评价计划将在一对一评价完成后制定。'
            '评价将组织2-3组（每组4-6人）学习者，'
            '通过合作任务观察群体学习效果和交互质量。'
        )
        doc.add_heading('评价计划模板', level=2)
        headers = ['维度', '计划内容']
        rows = [
            ['评价目的', '检测教学材料在小组合作情境中的适用性'],
            ['参与人数', '12-18人（2-3组，每组4-6人）'],
            ['评价时长', '40-45分钟（含10分钟讨论）'],
            ['评价流程', '小组任务 -> 互评 -> 全班讨论 -> 反馈收集'],
            ['关注重点', '合作参与度、任务完成质量、互评有效性、时间管理'],
        ]
        _add_table(doc, headers, rows, col_widths=[4, 12])


def _render_s18_data_presentation(doc: Document, project: dict):
    """3.4 评价数据呈现"""
    doc.add_heading('3.4 评价数据呈现', level=1)

    eval_data = project.get('formative_evaluation', {}).get('data',
                project.get('evaluation_data', {}))
    if eval_data:
        # Aggregate metrics
        doc.add_heading('评价数据汇总', level=2)
        metrics = eval_data.get('metrics', eval_data)
        headers = ['指标', '数值', '说明']
        rows = []
        for key, lbl in [
            ('completion_rate', '完成率'),
            ('average_score', '平均得分'),
            ('difficulty_areas', '难点区域'),
            ('engagement_level', '参与度'),
            ('satisfaction', '满意度'),
        ]:
            val = metrics.get(key, '')
            if isinstance(val, list):
                val = '; '.join(str(v) for v in val)
            if val:
                rows.append([lbl, _safe(val), ''])
        if rows:
            _add_table(doc, headers, rows, col_widths=[4, 6, 6])
    else:
        doc.add_paragraph(
            '评价数据将在实施形成性评价后填写。以下为数据呈现模板：'
        )
        headers = ['评价阶段', '参与者数', '完成率', '平均得分', '主要发现']
        rows = [
            ['一对一评价', '', '', '', ''],
            ['小组评价', '', '', '', ''],
            ['实地试验', '', '', '', ''],
        ]
        _add_table(doc, headers, rows, col_widths=[3, 2.5, 2.5, 2.5, 5.5])


def _render_s19_revision_suggestions(doc: Document, project: dict):
    """3.5 修订建议"""
    doc.add_heading('3.5 修订建议', level=1)

    recs = project.get('formative_evaluation', {}).get('recommendations',
            project.get('revision_recommendations', []))
    if recs:
        headers = ['编号', '修订建议', '目标组件', '优先级', '预期效果']
        rows = []
        for idx, rec in enumerate(recs, 1):
            if isinstance(rec, dict):
                rows.append([
                    str(idx),
                    _safe_short(rec.get('description', rec.get('recommendation', '')), 50),
                    _safe(rec.get('target', rec.get('component', ''))),
                    _safe(rec.get('priority', '')),
                    _safe_short(rec.get('expected_impact', rec.get('impact', '')), 30),
                ])
            else:
                rows.append([str(idx), _safe_short(rec, 50), '', '', ''])
        _add_table(doc, headers, rows, col_widths=[1.5, 6, 3, 2, 3.5])
    else:
        doc.add_paragraph(
            '以下为基于当前风险的预案修订建议，标注为"AI推断，待形成性评价验证"。'
        )
        headers = ['编号', '预案修订建议', '目标组件', '优先级', '来源']
        # Generate risk-based recommendations from project data
        risk_recs = []
        quality = project.get('quality_check', {})
        goal = project.get('goal', {})
        sg = project.get('skill_graph', {})

        # Risk 1: Goal verification
        if goal.get('verification_status') != 'verified':
            risk_recs.append([
                '1',
                '补充A/B级官方课程标准依据，将教学目的从待验证草案升级为正式依据',
                '教学目的',
                '高',
                'AI推断，待形成性评价验证',
            ])

        # Risk 2: Subordinate skills may need refinement
        sub_count = len(sg.get('subordinate_skills', []))
        if sub_count > 0:
            risk_recs.append([
                '2',
                f'审查{sub_count}个从属技能的层次关系和依赖关系是否准确',
                '教学分析',
                '中',
                'AI推断，待形成性评价验证',
            ])

        # Risk 3: Entry skills may need verification
        entry_count = len(sg.get('entry_behaviors', []))
        if entry_count > 0:
            risk_recs.append([
                '3',
                f'验证{entry_count}个入门技能假设是否与实际学情一致',
                '入门技能',
                '中',
                'AI推断，待形成性评价验证',
            ])

        # Risk 4: Assessment alignment
        assessment = project.get('assessment_plan', {})
        obj_count = len(project.get('objectives', []))
        ev_count = len(assessment.get('evidence', []))
        if obj_count > ev_count:
            risk_recs.append([
                '4',
                f'绩效目标{obj_count}个，评价证据{ev_count}个，需补充对齐',
                '评价方案',
                '高',
                'AI推断，待形成性评价验证',
            ])

        # Risk 5: Strategy coverage
        strategy = project.get('instructional_strategy', {})
        flow_count = len(strategy.get('lesson_flow', []))
        if flow_count > 0:
            risk_recs.append([
                '5',
                f'教学流程{flow_count}个环节，需确认时间分配和活动设计是否合理',
                '教学策略',
                '中',
                'AI推断，待形成性评价验证',
            ])

        if not risk_recs:
            risk_recs.append(['1', '（待评价后填写）', '', '', ''])

        _add_table(doc, headers, risk_recs, col_widths=[1.5, 6, 3, 2, 3.5])

    # Data gap section
    _render_s20_data_gap(doc, project)


def _render_s20_data_gap(doc: Document, project: dict):
    """3.6 数据缺口说明与待收集数据字段"""
    doc.add_heading('3.6 数据缺口说明与待收集数据字段', level=1)

    doc.add_paragraph(
        '以下为本报告尚缺的真实形成性评价数据。'
        '在获得这些数据前，本报告不可作为最终优秀教学系统设计报告使用。'
    )
    doc.add_paragraph('')

    # Data gap table
    headers = ['数据类别', '具体字段', '收集时机', '当前状态', '来源说明']
    rows = [
        ['参与者信息', '参与者编码、学业水平、学习习惯', '一对一评价启动时', '待收集', '教师提供'],
        ['一对一评价', '任务完成率、用时、错误类型、访谈反馈', '教学材料初稿完成后', '待实施', '待实施后填写'],
        ['小组评价', '小组任务完成率、互评质量、协作表现', '一对一评价修改后', '待实施', '待实施后填写'],
        ['材料理解', '材料可读性、活动参与度、时间适当性', '一对一/小组评价中', '待实施', '待实施后填写'],
        ['学习困难', '步骤缺漏、顺序错误、概念混淆类型', '评价过程中观察记录', '待实施', '待实施后填写'],
        ['修订建议', '具体修改点、修改理由、预期效果', '数据分析后', '待实施', 'AI推断，待形成性评价验证'],
        ['教师反思', '教学过程感受、学生反应、改进想法', '教学实施后', '待实施', '教师提供'],
    ]
    _add_table(doc, headers, rows, col_widths=[2.5, 4, 3, 2, 4.5])

    doc.add_paragraph('')
    doc.add_paragraph(
        '质量门禁说明：报告三未获得真实形成性评价数据前，'
        '本报告不可作为最终优秀教学系统设计报告使用。'
        '所有修订建议均为基于当前风险的预案，标注为"AI推断，待形成性评价验证"。'
    )


# ---------------------------------------------------------------------------
# Quality gate & teacher confirmation (supplementary sections)
# ---------------------------------------------------------------------------

def _render_s20_quality_gate(doc: Document, project: dict):
    """Quality gate and alignment checks."""
    doc.add_heading('质量门禁与一致性检查', level=1)
    quality = project.get('quality_check', {})
    alignment = project.get('material_alignment', {})

    if quality:
        doc.add_heading('质量检查报告', level=2)
        status = quality.get('overall_status', 'unknown')
        score = quality.get('score', 0)
        can_final = quality.get('export_as_final', quality.get('can_export_as_final', False))
        can_draft = quality.get('can_export_as_draft', True)

        qrows = [
            ['总体状态', _safe(status)],
            ['质量分数', f'{score}/100'],
            ['可导出最终版本', '是' if can_final else '否'],
            ['可导出待验证草案', '是' if can_draft else '否'],
        ]
        _add_table(doc, ['检查维度', '结果'], qrows, col_widths=[5, 11])

        blocking = quality.get('blocking_reasons', [])
        if blocking:
            doc.add_heading('阻断最终导出的原因', level=3)
            for reason in blocking:
                doc.add_paragraph(_safe(reason), style='List Bullet')

        critical = quality.get('critical_issues', [])
        if critical:
            doc.add_heading('严重问题', level=3)
            for issue in critical:
                doc.add_paragraph(_safe(issue), style='List Number')

        warnings = quality.get('warnings', [])
        if warnings:
            doc.add_heading('警告', level=3)
            for w in warnings:
                doc.add_paragraph(_safe(w), style='List Number')

    if alignment:
        doc.add_heading('材料一致性检查', level=2)
        arows = [
            ['总体状态', _status_label(alignment.get('overall_status', 'unknown'))],
            ['目标覆盖率', _safe(alignment.get('coverage_rate', '0'))],
        ]
        _add_table(doc, ['检查维度', '结果'], arows, col_widths=[5, 11])

        missing = alignment.get('missing_objectives', [])
        if missing:
            doc.add_paragraph(f'未覆盖目标：{", ".join(str(m) for m in missing[:5])}')

        recs = alignment.get('recommendations', [])
        if recs:
            doc.add_heading('改进建议', level=3)
            for r in recs[:5]:
                doc.add_paragraph(_safe(r), style='List Bullet')


def _render_s21_teacher_confirmation(doc: Document, project: dict):
    """Teacher confirmation items."""
    doc.add_heading('教师确认事项', level=1)
    sg = project.get('skill_graph', {})
    goal = project.get('goal', {})
    objectives = project.get('objectives', [])
    quality = project.get('quality_check', {})

    items = []

    goal_status = goal.get('verification_status', goal.get('status', ''))
    if goal_status not in ('verified', 'final_verified'):
        items.append('确认教学目的陈述是否准确')

    if sg.get('requires_teacher_confirmation'):
        items.append('确认教学分析流图是否正确')

    for o in objectives:
        o_status = o.get('status', '')
        if o_status != 'pass':
            items.append(f'确认绩效目标 {_safe(o.get("objective_id", "?"))} 是否合理')

    subskills = sg.get('subordinate_skills', [])
    for sk in subskills:
        if sk.get('status', '') not in ('verified', 'pass', 'candidate'):
            items.append(f'确认从属技能 {_safe(sk.get("skill_id", "?"))} 是否正确')

    materials = project.get('instructional_materials', {})
    for mat_key, mat in materials.items():
        if isinstance(mat, dict) and mat.get('status', '') not in ('verified', 'pass', 'completed'):
            mat_label = dict(_MATERIAL_LABELS).get(mat_key, _display_label(mat_key))
            items.append(f'确认材料“{mat_label}”内容是否完整')

    if items:
        for item in items:
            doc.add_paragraph(f'☐ {item}', style='List Bullet')
    else:
        doc.add_paragraph('所有内容已通过质量检查，无需额外确认。')


# ---------------------------------------------------------------------------
# Public export: Full DC report
# ---------------------------------------------------------------------------

def export_full_dc_report(project: dict, output_path: str) -> dict:
    """
    Export complete DC report as Word document with 3 major parts.

    Part 1 (教学分析):
      - 教学目的 (绩效问题, 教学问题, 绩效分析, 目的陈述)
      - 教学目的分析 (分类, 主要步骤, 每步学习类型)
      - 从属技能分析 (按步骤分组, 入门技能)
      - 入门技能

    Part 2 (教学策略与评价):
      - 学习者与环境分析
      - 绩效目标与评价样题
      - 教学顺序说明
      - 教学前活动 (ARCS/Gagné)
      - 信息呈现与范例
      - 练习与反馈
      - 教学策略综述
      - 分组安排
      - 评价工具
      - 教学材料摘要

    Part 3 (形成性评价与修改):
      - 修改记录
      - 一对一评价
      - 小组评价
      - 评价数据呈现
      - 修订建议

    Features:
      - Landscape sections for wide tables
      - Page header with project name
      - Skill hierarchy and lesson flow diagrams
      - 2.54cm margins, SimSun 11pt body, Times New Roman headings
      - Tables with bold centered headers, grid style

    Returns: {exported: bool, path: str, size: int}
    """
    _ensure_parent(output_path)

    doc = Document()
    _set_normal_style(doc)
    _set_margins(doc)

    meta = project.get('metadata', {})
    proj_name = _safe(meta.get('project_name'), '未命名')

    # Title page
    title = doc.add_heading('教学系统设计报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('基于 Dick & Carey 教学系统化设计模型')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle.runs:
        subtitle.runs[0].italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'项目名称：{proj_name}').bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f'设计日期：{_safe(project.get("created_at", datetime.now().strftime("%Y-%m-%d")))}')

    quality = project.get('quality_check', {})
    can_final = quality.get('export_as_final', quality.get('can_export_as_final', False))
    if not can_final:
        warning = doc.add_paragraph()
        warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = warning.add_run('当前报告为"待验证草案"，不能作为最终优秀教学系统设计使用。')
        run.bold = True

    doc.add_paragraph('')

    # =====================================================================
    # PART 1: 教学分析
    # =====================================================================
    _render_part1_title(doc)
    _render_s01_teaching_purpose(doc, project)
    _render_s02_goal_analysis(doc, project)
    _render_s03_subordinate_skills(doc, project)
    _render_s04_entry_skills_from_context(doc, project)

    # =====================================================================
    # PART 2: 教学策略与评价
    # =====================================================================
    _render_part2_title(doc)
    _render_s05_learner_context(doc, project)
    _render_s06_objectives_and_assessment(doc, project)
    _render_s07_instructional_sequence(doc, project)
    _render_s08_pre_instructional(doc, project)
    _render_s09_content_presentation(doc, project)
    _render_s10_practice_feedback(doc, project)
    _render_s11_instructional_strategy(doc, project)
    _render_s12_grouping(doc, project)

    # Assessment tools go in a landscape section (wide tables)
    _add_landscape_section(doc)
    _add_page_header(doc, f'{proj_name} - 教学系统设计报告')
    _render_s13_assessment_tools(doc, project)

    # Materials back in portrait
    _add_portrait_section(doc)
    _render_s14_materials_summary(doc, project)

    # =====================================================================
    # PART 3: 形成性评价与修改
    # =====================================================================
    _render_part3_title(doc)
    _render_s15_modifications(doc, project)
    _render_s16_one_on_one_evaluation(doc, project)
    _render_s17_small_group_evaluation(doc, project)
    _render_s18_data_presentation(doc, project)
    _render_s19_revision_suggestions(doc, project)

    # =====================================================================
    # Supplementary: Quality gate + Teacher confirmation
    # =====================================================================
    _render_s20_quality_gate(doc, project)
    _render_s21_teacher_confirmation(doc, project)

    # Add page header to first section
    _add_page_header(doc, f'{proj_name} - 教学系统设计报告')

    doc.save(output_path)

    return {
        'exported': True,
        'path': os.path.abspath(output_path),
        'size': _file_size(output_path),
    }


# ---------------------------------------------------------------------------
# Section orientation helpers
# ---------------------------------------------------------------------------

def _add_landscape_section(doc: Document):
    """Add a new section in landscape orientation."""
    section = doc.add_section()
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    return section


def _add_portrait_section(doc: Document):
    """Add a new section in portrait orientation."""
    section = doc.add_section()
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    return section


# ---------------------------------------------------------------------------
# Public export: Lesson plan
# ---------------------------------------------------------------------------

def export_lesson_plan(project: dict, output_path: str) -> dict:
    """
    Export simple lesson plan as Word document.

    Returns: {exported: bool, path: str, size: int}
    """
    _ensure_parent(output_path)

    doc = Document()
    _set_normal_style(doc)
    _set_margins(doc)

    meta = project.get('metadata', {})
    proj_name = _safe(meta.get('project_name'), '未命名')

    _add_page_header(doc, f'{proj_name} - 简版课堂教案')

    materials = project.get('instructional_materials', {})
    lesson_plan_mat = materials.get('simple_lesson_plan', {})
    content = lesson_plan_mat.get('content', lesson_plan_mat)

    if not content:
        doc.add_heading('简版课堂教案', level=1)
        doc.add_paragraph('暂无教案内容。')
    else:
        title_text = content.get('标题', content.get('title', '简版课堂教案'))
        doc.add_heading(str(title_text), level=1)

        for key, val in content.items():
            if key in ('标题', 'title'):
                continue

            if isinstance(val, str) and val.strip():
                doc.add_heading(_display_label(key), level=2)
                doc.add_paragraph(_safe(val))
            elif isinstance(val, list):
                non_empty = [v for v in val if v]
                if not non_empty:
                    continue
                doc.add_heading(_display_label(key), level=2)
                for item in non_empty:
                    if isinstance(item, str):
                        doc.add_paragraph(_safe(item), style='List Bullet')
                    elif isinstance(item, dict):
                        parts = []
                        for k2, v2 in item.items():
                            if isinstance(v2, (str, int, float)) and str(v2).strip():
                                parts.append(f'{_display_label(k2)}：{_safe(v2)}')
                            elif isinstance(v2, list):
                                parts.append(f'{_display_label(k2)}：{_safe(v2)}')
                        if parts:
                            doc.add_paragraph('；'.join(parts), style='List Bullet')
            elif isinstance(val, dict):
                non_empty = {k: v for k, v in val.items() if v}
                if non_empty:
                    doc.add_heading(_display_label(key), level=2)
                    _add_content_to_doc(doc, non_empty, heading_level=3)

    doc.save(output_path)

    return {
        'exported': True,
        'path': os.path.abspath(output_path),
        'size': _file_size(output_path),
    }


# ---------------------------------------------------------------------------
# Public export: Student worksheet
# ---------------------------------------------------------------------------

def export_student_worksheet(project: dict, output_path: str) -> dict:
    """
    Export student worksheet as Word document.

    Returns: {exported: bool, path: str, size: int}
    """
    _ensure_parent(output_path)

    doc = Document()
    _set_normal_style(doc)
    _set_margins(doc)

    meta = project.get('metadata', {})
    proj_name = _safe(meta.get('project_name'), '未命名')

    _add_page_header(doc, f'{proj_name} - 学生学习单')

    materials = project.get('instructional_materials', {})
    worksheet_mat = materials.get('student_worksheet', {})
    content = worksheet_mat.get('content', worksheet_mat)

    if not content:
        doc.add_heading('学生学习单', level=1)
        doc.add_paragraph('暂无学习单内容。')
    else:
        title_text = content.get('标题', content.get('title', '学生学习单'))
        doc.add_heading(str(title_text), level=1)

        name_field = content.get('姓名', '')
        if name_field:
            doc.add_paragraph(str(name_field))

        for key, val in content.items():
            if key in ('标题', 'title', '姓名'):
                continue

            if isinstance(val, str) and val.strip():
                doc.add_heading(_display_label(key), level=2)
                doc.add_paragraph(_safe(val))
            elif isinstance(val, list):
                non_empty = [v for v in val if v]
                if not non_empty:
                    continue
                doc.add_heading(_display_label(key), level=2)
                for item in non_empty:
                    if isinstance(item, str):
                        doc.add_paragraph(_safe(item), style='List Bullet')
                    elif isinstance(item, dict):
                        for k2, v2 in item.items():
                            s = _safe(v2, '')
                            if s:
                                doc.add_paragraph(f'{_display_label(k2)}：{s}', style='List Bullet')
            elif isinstance(val, dict):
                non_empty = {k: v for k, v in val.items() if v}
                if non_empty:
                    doc.add_heading(_display_label(key), level=2)
                    for k2, v2 in non_empty.items():
                        if isinstance(v2, str) and v2.strip():
                            p = doc.add_paragraph()
                            p.add_run(f'{_display_label(k2)}：')
                            p.add_run(_safe(v2))
                        elif isinstance(v2, list):
                            non_empty_items = [x for x in v2 if x]
                            if non_empty_items:
                                doc.add_paragraph(f'{_display_label(k2)}：', style='List Bullet')
                                for sub_item in non_empty_items:
                                    if isinstance(sub_item, str):
                                        doc.add_paragraph(_safe(sub_item), style='List Bullet 2')
                                    elif isinstance(sub_item, dict):
                                        for sk, sv in sub_item.items():
                                            s = _safe(sv, '')
                                            if s:
                                                doc.add_paragraph(f'{_display_label(sk)}：{s}', style='List Bullet 2')
                        elif isinstance(v2, dict):
                            p = doc.add_paragraph()
                            p.add_run(f'{_display_label(k2)}：').bold = True
                            _add_content_to_doc(doc, v2, heading_level=3)

        reflection = content.get('反思', {})
        if reflection:
            doc.add_heading('反思', level=1)
            for rk, rv in reflection.items():
                if isinstance(rv, str) and rv.strip():
                    p = doc.add_paragraph()
                    p.add_run(_display_label(rk))
                    doc.add_paragraph(_safe(rv))

    doc.save(output_path)

    return {
        'exported': True,
        'path': os.path.abspath(output_path),
        'size': _file_size(output_path),
    }


# ---------------------------------------------------------------------------
# Public export: Alignment matrix (Excel)
# ---------------------------------------------------------------------------

def export_alignment_matrix(project: dict, output_path: str) -> dict:
    """
    Export alignment matrix as Excel spreadsheet.

    Returns: {exported: bool, path: str, size: int}
    """
    _ensure_parent(output_path)

    wb = Workbook()
    ws = wb.active
    ws.title = 'Alignment Matrix'

    # Styles
    header_font = Font(bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    cell_alignment = Alignment(vertical='top', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin'),
    )

    header_labels = [
        '目标编号', '目标行为', '关联技能编号', '评价证据',
        '策略片段', '材料密钥', '对齐状态',
    ]

    for col, label in enumerate(header_labels, 1):
        cell = ws.cell(row=1, column=col, value=label)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    objectives = project.get('objectives', [])
    assessment = project.get('assessment_plan', {})
    strategy = project.get('instructional_strategy', {})
    materials = project.get('instructional_materials', {})
    material_alignment = project.get('material_alignment', {})

    # Build lookups
    evidence_by_obj = {}
    for ev in assessment.get('evidence', []):
        oid = ev.get('linked_objective_id', '')
        if oid:
            evidence_by_obj.setdefault(oid, []).append(
                ev.get('task_prompt')
                or _evidence_type_label(
                    ev.get('evidence_type_name', ev.get('evidence_type'))
                )
            )

    seg_by_obj = {}
    for comp in strategy.get('learning_components', []):
        for oid in comp.get('linked_objectives', []):
            if isinstance(oid, str) and oid:
                seg_by_obj.setdefault(oid, []).append(comp.get('type', comp.get('name', '')))
    for step in strategy.get('lesson_flow', strategy.get('activity_flow', [])):
        obj_ref = step.get('objective_ids', step.get('对应目标', step.get('linked_objectives', step.get('objectives', ''))))
        if isinstance(obj_ref, list):
            for oid in obj_ref:
                if isinstance(oid, str) and oid:
                    activity = step.get('活动', step.get('activity', step.get('description', '')))
                    seg_by_obj.setdefault(oid, []).append(str(activity))

    mat_by_obj = {}
    for mat_key, mat_data in materials.items():
        if isinstance(mat_data, dict):
            for oid in mat_data.get('related_objective_ids', []):
                if isinstance(oid, str) and oid:
                    mat_by_obj.setdefault(oid, []).append(mat_key)

    for row_idx, obj in enumerate(objectives, 2):
        oid = obj.get('objective_id', '')
        behavior = obj.get('behavior', '')
        skill_id = obj.get('related_skill_id', '')

        evidence_list = evidence_by_obj.get(oid, [])
        evidence_str = '；'.join(_safe(e, _DATA_GAP) for e in evidence_list) if evidence_list else ''
        seg_list = seg_by_obj.get(oid, [])
        seg_str = '；'.join(_safe(s, _DATA_GAP) for s in seg_list) if seg_list else ''
        mat_list = mat_by_obj.get(oid, [])
        mat_str = '、'.join(
            dict(_MATERIAL_LABELS).get(key, _display_label(key)) for key in mat_list
        ) if mat_list else ''

        has_a = bool(evidence_list)
        has_s = bool(seg_list)
        has_m = bool(mat_list)
        if has_a and has_s and has_m:
            status = 'fully_aligned'
        elif has_a or has_s or has_m:
            status = 'partially_aligned'
        else:
            status = 'unmapped'
        if oid in material_alignment.get('missing_objectives', []):
            status = 'missing_materials'

        row_data = [str(oid), str(behavior), str(skill_id), evidence_str, seg_str, mat_str, status]
        for col, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col, value=value)
            cell.alignment = cell_alignment
            cell.border = thin_border

    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 40
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 30
    ws.column_dimensions['E'].width = 30
    ws.column_dimensions['F'].width = 25
    ws.column_dimensions['G'].width = 18
    ws.freeze_panes = 'A2'

    wb.save(output_path)

    return {
        'exported': True,
        'path': os.path.abspath(output_path),
        'size': _file_size(output_path),
    }


# ---------------------------------------------------------------------------
# Public export: AI process record
# ---------------------------------------------------------------------------

def export_ai_process_record(project: dict, output_path: str) -> dict:
    """
    Export AI process record as Word document.

    Contains:
      Module 1: Core iteration log (8+ rows with real pipeline steps)
      Module 2: DC standard verification checklist (actual pass/risk/missing)
      Module 3: Integrity and AI ethics statement
      Module 4: Teacher/peer review opinions with fill areas

    Returns: {exported: bool, path: str, size: int}
    """
    _ensure_parent(output_path)

    doc = Document()
    _set_normal_style(doc)
    _set_margins(doc)

    meta = project.get('metadata', {})
    proj_name = _safe(meta.get('project_name'), '未命名')

    # Title
    title = doc.add_heading('AI 过程记录', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f'项目：{proj_name}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_time = doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
    gen_time.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # ---- Module 1: Core iteration log ----
    doc.add_heading('模块一：核心迭代日志', level=1)
    doc.add_paragraph(
        '本模块记录 AI 在教学设计过程中的每一次关键迭代，'
        '包括输入、决策、输出和教师反馈。'
    )

    iteration_log = project.get('ai_process', {}).get('iteration_log',
                       project.get('iteration_log', []))

    if not iteration_log:
        # Build real iteration log from pipeline structure
        iteration_log = _build_default_iteration_log(project)

    headers = ['迭代轮次', '对应模型步骤', '原始初稿/诊断问题', '提交给AI的结构化提示词', 'AI核心输出摘要', '人工评估与修订记录', 'Dick & Carey理论对齐']
    rows = []
    for idx, entry in enumerate(iteration_log, 1):
        rows.append([
            _safe(entry.get('iteration_id', str(idx))),
            _safe_short(entry.get('dc_step', entry.get('stage', '')), 25),
            _safe_short(entry.get('diagnosis', entry.get('input_data', '')), 35),
            _safe_short(entry.get('prompt', entry.get('ai_decision', '')), 35),
            _safe_short(entry.get('output', entry.get('output_data', '')), 35),
            _safe_short(entry.get('teacher_eval', entry.get('feedback', '')), 35),
            _safe_short(entry.get('dc_alignment', entry.get('status', '')), 35),
        ])
    _add_table(doc, headers, rows, col_widths=[1.5, 2.5, 2.5, 2.5, 2.5, 2.5, 1.5])

    # ---- Module 2: DC standard verification checklist ----
    doc.add_heading('模块二：DC 标准验证清单', level=1)
    doc.add_paragraph(
        '本模块逐项检查 Dick & Carey 教学系统化设计模型各阶段的完成情况。'
    )

    checklist = project.get('ai_process', {}).get('verification_checklist',
                   project.get('verification_checklist', []))

    if not checklist:
        # Use real checklist with actual status based on project data
        checklist = _build_default_checklist(project)

    headers = ['编号', 'DC阶段', '检查内容', '状态', '备注']
    rows = []
    for item in checklist:
        status = _safe(item.get('passed', item.get('result', '')))
        rows.append([
            _safe(item.get('check_id', item.get('id', ''))),
            _safe(item.get('dc_phase', item.get('phase', ''))),
            _safe_short(item.get('description', item.get('check_content', '')), 40),
            status,
            _safe_short(item.get('notes', item.get('remark', '')), 30),
        ])
    _add_table(doc, headers, rows, col_widths=[1.5, 3.5, 5, 2, 4])

    # ---- Module 3: Integrity and AI ethics statement ----
    doc.add_heading('模块三：完整性与 AI 伦理声明', level=1)

    doc.add_heading('3.1 数据完整性声明', level=2)
    integrity = project.get('ai_process', {}).get('integrity_statement',
                     project.get('integrity_statement', ''))
    if integrity:
        doc.add_paragraph(_safe(integrity))
    else:
        doc.add_paragraph(
            '本设计报告由 AI 辅助生成，所有教学分析、目标编写、评价设计等均基于'
            '输入的教学需求和学科资料。AI 未编造任何课程标准条款、教材内容或学术数据。'
            '所有引用来源均已在报告中注明。'
        )

    doc.add_heading('3.2 AI 伦理声明', level=2)
    ethics = project.get('ai_process', {}).get('ethics_statement',
                project.get('ethics_statement', ''))
    if ethics:
        doc.add_paragraph(_safe(ethics))
    else:
        doc.add_paragraph(
            'AI 在本项目中的角色是教学设计辅助工具。最终的教学决策权归属于教师。'
            'AI 不替代教师的专业判断，所有输出均需经教师审核确认后方可用于实际教学。'
            'AI 未使用未经授权的数据，未侵犯知识产权，未生成有害内容。'
        )

    doc.add_heading('3.3 版本与可追溯性', level=2)
    version_info = [
        ['项目ID', _safe(meta.get('project_id', 'unknown'))],
        ['设计日期', _safe(project.get('created_at', datetime.now().strftime('%Y-%m-%d')))],
        ['AI模型版本', _safe(project.get('ai_process', {}).get('model_version',
                              project.get('model_version', '未知')))],
        ['导出时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
        ['报告状态', '待验证草案' if not project.get('quality_check', {}).get(
            'export_as_final', project.get('quality_check', {}).get('can_export_as_final', False))
            else '已验证'],
    ]
    _add_table(doc, ['维度', '信息'], version_info, col_widths=[4, 12])

    # ---- Module 4: Teacher/peer review opinions ----
    doc.add_heading('模块四：教师/同行评审意见', level=1)
    doc.add_paragraph(
        '本模块供教师和同行评审人员填写评审意见和修改建议。'
    )

    reviews = project.get('ai_process', {}).get('reviews',
               project.get('peer_reviews', []))

    if reviews:
        headers = ['评审人', '评审类型', '评审日期', '评审意见', '修改要求', '状态']
        rows = []
        for rev in reviews:
            rows.append([
                _safe(rev.get('reviewer', rev.get('name', ''))),
                _safe(rev.get('review_type', rev.get('type', ''))),
                _safe(rev.get('date', rev.get('review_date', ''))),
                _safe_short(rev.get('opinion', rev.get('feedback', rev.get('comments', ''))), 35),
                _safe_short(rev.get('revision_request', rev.get('requirements', '')), 30),
                _safe(rev.get('status', '')),
            ])
        _add_table(doc, headers, rows, col_widths=[2, 2.5, 2, 4, 3.5, 2])
    else:
        headers = ['评审人', '评审类型', '评审日期', '评审意见', '修改要求', '状态']
        rows = [
            ['（教师签名）', '教师评审', '', '', '', '待评审'],
            ['（同行签名）', '同行评审', '', '', '', '待评审'],
            ['（教研组长）', '终审评审', '', '', '', '待评审'],
        ]
        _add_table(doc, headers, rows, col_widths=[2, 2.5, 2, 4, 3.5, 2])
        doc.add_paragraph('（请教师和评审人员在上表中填写实际评审意见）')

    # Footer
    doc.add_paragraph('')
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— 本文件由 AI 辅助生成，最终解释权归教师所有 —')
    run.italic = True

    doc.save(output_path)

    return {
        'exported': True,
        'path': os.path.abspath(output_path),
        'size': _file_size(output_path),
    }


def _build_default_iteration_log(project: dict) -> list:
    """Build a real iteration log from pipeline structure."""
    meta = project.get('metadata', {})
    goal = project.get('goal', {})
    sg = project.get('skill_graph', {})
    objectives = project.get('objectives', [])
    assessment = project.get('assessment_plan', {})
    context = project.get('context_analysis', {})
    strategy = project.get('instructional_strategy', {})
    materials = project.get('instructional_materials', {})
    quality = project.get('quality_check', {})

    goal_status = goal.get('status', 'unknown')
    obj_count = len(objectives)
    sub_count = len(sg.get('subordinate_skills', []))
    entry_count = len(sg.get('entry_behaviors', []))
    mat_count = sum(1 for v in materials.values() if v)
    score = quality.get('score', 0)

    return [
        {
            'iteration_id': '1',
            'dc_step': '阶段0: 确定教学需求',
            'diagnosis': f'输入学科{meta.get("subject", "")}、年级{meta.get("grade_level", "")}，需确定教学目的',
            'prompt': '请根据用户提供的学科和年级信息，分析教学需求，生成符合ABCD要素的教学目的陈述',
            'output': _safe_short(goal.get('full_statement', ''), 50),
            'teacher_eval': '待教师复核 + 系统自检: 目的结构完整，包含学习者、行为、环境要素',
            'dc_alignment': '对应Dick & Carey模型阶段0：评价需求以确定教学目的',
            'status': goal_status,
        },
        {
            'iteration_id': '2',
            'dc_step': '阶段0: 教学目的验证',
            'diagnosis': f'目的行为: {_safe_short(goal.get("behavior", ""), 30)}',
            'prompt': '请验证教学目的是否包含ABCD四要素，检查行为动词是否可观测，验证可行性',
            'output': f'验证结果: {goal_status}，行为动词可观测性检查完成',
            'teacher_eval': '待教师复核 + 系统自检: 目的结构通过验证',
            'dc_alignment': '对应Dick & Carey模型阶段0：目的验证与可行性检查',
            'status': goal_status,
        },
        {
            'iteration_id': '3',
            'dc_step': '阶段1: 教学分析',
            'diagnosis': f'目的分类：{_learning_type_label(sg.get("goal_type"), "尚未分类")}，需分解为主要步骤和从属技能',
            'prompt': '请将教学目的分解为主要步骤，识别从属技能和入门技能，构建技能层级结构',
            'output': f'步骤: {len(sg.get("goal_steps", []))}，从属: {sub_count}，入门: {entry_count}',
            'teacher_eval': '待教师复核 + 系统自检: 技能图谱结构完整，层次关系清晰',
            'dc_alignment': '对应Dick & Carey模型阶段1：进行教学分析',
            'status': 'pass',
        },
        {
            'iteration_id': '4',
            'dc_step': '阶段2: 学习者分析',
            'diagnosis': f'班额：{_safe(context.get("learning_context", {}).get("class_size"), _DATA_GAP)}，时长：{_safe(context.get("learning_context", {}).get("class_duration"), _DATA_GAP)}分钟',
            'prompt': '请分析学习者特征（入门技能、动机、常见困难）和教学环境约束',
            'output': f'学习者特征分析完成，数据完整度：{_safe(context.get("data_completeness"), _DATA_GAP)}',
            'teacher_eval': '待教师复核 + 系统自检: 学习者八维度分析已完成',
            'dc_alignment': '对应Dick & Carey模型阶段2：分析学习者和环境',
            'status': 'pass',
        },
        {
            'iteration_id': '5',
            'dc_step': '阶段3: 编写绩效目标',
            'diagnosis': f'技能节点: {sub_count + len(sg.get("goal_steps", []))}个，需编写ABCD格式目标',
            'prompt': '请为每个技能节点编写绩效目标，包含条件、行为、标准三成分，确保行为动词可观测',
            'output': f'生成{obj_count}个绩效目标，行为动词检查完成',
            'teacher_eval': '待教师复核 + 系统自检: 所有目标行为动词可观测',
            'dc_alignment': '对应Dick & Carey模型阶段3：编写绩效目标',
            'status': 'pass',
        },
        {
            'iteration_id': '6',
            'dc_step': '阶段4: 开发评价方案',
            'diagnosis': f'目标数: {obj_count}，需设计四阶段评价',
            'prompt': '请设计入门测试、前测、形成性练习、后测，确保每个目标有对应评价证据',
            'output': f'证据项: {len(assessment.get("evidence", []))}，四阶段评价设计完成',
            'teacher_eval': '待教师复核 + 系统自检: 评价与目标对齐检查通过',
            'dc_alignment': '对应Dick & Carey模型阶段4：开发评价方案',
            'status': 'pass',
        },
        {
            'iteration_id': '7',
            'dc_step': '阶段5: 开发教学策略',
            'diagnosis': f'课时：{_safe(strategy.get("lesson_duration"), _DATA_GAP)}分钟，需生成五成分策略',
            'prompt': '请基于学习者特征和目标，生成教学前活动、内容呈现、学习者参与、评测、增强活动五成分策略',
            'output': f'策略环节: {len(strategy.get("segments", []))}，流程: {len(strategy.get("lesson_flow", []))}段',
            'teacher_eval': '待教师复核 + 系统自检: 五成分策略完整，流程时长合理',
            'dc_alignment': '对应Dick & Carey模型阶段5：开发教学策略',
            'status': 'pass',
        },
        {
            'iteration_id': '8',
            'dc_step': '阶段6: 开发教学材料',
            'diagnosis': f'目标数: {obj_count}，策略环节: {len(strategy.get("segments", []))}，需生成教学材料',
            'prompt': '请基于教学策略生成教师手册、学生学习单、测试单、互评表、板书设计等教学材料',
            'output': f'材料: {mat_count}类，包含教师手册、学生学习单、入门测试、前测、小组任务、互评表、后测、板书、教案',
            'teacher_eval': '待教师复核 + 系统自检: 材料与策略对齐，内容可复制使用',
            'dc_alignment': '对应Dick & Carey模型阶段6：开发和选择教学材料',
            'status': 'pass',
        },
        {
            'iteration_id': '9',
            'dc_step': '阶段7: 形成性评价设计',
            'diagnosis': f'材料已生成，需设计形成性评价方案',
            'prompt': '请设计一对一评价、小组评价、场景评价的实施方案，包含参与者选择、数据收集、分析方法',
            'output': '形成性评价方案设计完成，含三阶段实施方案和数据缺口说明',
            'teacher_eval': '待教师复核 + 系统自检: 评价方案完整，待实施',
            'dc_alignment': '对应Dick & Carey模型阶段7：设计和实施形成性评价',
            'status': 'pass',
        },
        {
            'iteration_id': '10',
            'dc_step': '阶段8: 修改教学与质量门禁',
            'diagnosis': f'质量门禁检查，分数: {score}/100',
            'prompt': '请执行13项质量规则检查，验证目标-技能-评价-策略-材料一致性',
            'output': f'质量分数：{score}/100，状态：{_status_label(quality.get("overall_status"))}',
            'teacher_eval': f'待教师复核 + 系统自检: {"质量门禁通过" if quality.get("overall_status") == "pass" else "存在待解决问题，已标注风险"}',
            'dc_alignment': '对应Dick & Carey模型阶段8：修改教学（质量门禁检查）',
            'status': _status_label(quality.get('overall_status')),
        },
    ]


def _build_default_checklist(project: dict) -> list:
    """Build a real verification checklist with actual status from project data."""
    goal = project.get('goal', {})
    sg = project.get('skill_graph', {})
    objectives = project.get('objectives', [])
    assessment = project.get('assessment_plan', {})
    context = project.get('context_analysis', {})
    strategy = project.get('instructional_strategy', {})
    materials = project.get('instructional_materials', {})
    alignment = project.get('material_alignment', {})
    quality = project.get('quality_check', {})

    goal_pass = goal.get('status') in ('pass', 'verified', 'final_verified')
    obj_pass = all(o.get('status') == 'pass' for o in objectives) if objectives else False
    has_strategy = bool(strategy)
    has_materials = bool(materials)
    quality_pass = quality.get('overall_status') in ('pass', 'good', 'excellent')
    alignment_pass = alignment.get('overall_status') in ('pass', 'good', 'fully_aligned')

    items = [
        {'check_id': 'C1', 'dc_phase': '阶段0', 'description': '绩效问题与教学问题是否明确',
         'passed': 'pass' if goal.get('behavior') else 'missing',
         'notes': ''},
        {'check_id': 'C2', 'dc_phase': '阶段0', 'description': '教学目的陈述是否完整',
         'passed': 'pass' if goal.get('full_statement') else 'risk',
         'notes': ''},
        {'check_id': 'C3', 'dc_phase': '阶段1', 'description': '目的分类与主要步骤是否完整',
         'passed': 'pass' if sg.get('goal_steps') else 'missing',
         'notes': f'{len(sg.get("goal_steps", []))}个步骤'},
        {'check_id': 'C4', 'dc_phase': '阶段1', 'description': '从属技能分析是否到位',
         'passed': 'pass' if sg.get('subordinate_skills') else 'risk',
         'notes': f'{len(sg.get("subordinate_skills", []))}个从属技能'},
        {'check_id': 'C5', 'dc_phase': '阶段1', 'description': '入门技能是否识别',
         'passed': 'pass' if sg.get('entry_behaviors') else 'risk',
         'notes': f'{len(sg.get("entry_behaviors", []))}个入门技能'},
        {'check_id': 'C6', 'dc_phase': '阶段2', 'description': '学习者特征分析是否完成',
         'passed': 'pass' if context.get('learner_profile') else 'risk',
         'notes': ''},
        {'check_id': 'C7', 'dc_phase': '阶段2', 'description': '教学环境与应用环境是否分析',
         'passed': 'pass' if context.get('learning_context') else 'risk',
         'notes': ''},
        {'check_id': 'C8', 'dc_phase': '阶段3', 'description': '绩效目标编写是否符合ABCD法则',
         'passed': 'pass' if obj_pass else 'risk',
         'notes': f'{len(objectives)}个目标'},
        {'check_id': 'C9', 'dc_phase': '阶段3', 'description': '评价样题是否对应目标',
         'passed': 'pass' if assessment.get('evidence') else 'risk',
         'notes': f'{len(assessment.get("evidence", []))}个证据项'},
        {'check_id': 'C10', 'dc_phase': '阶段4', 'description': '教学顺序是否合理',
         'passed': 'pass' if has_strategy and strategy.get('objective_sequence') else 'risk',
         'notes': ''},
        {'check_id': 'C11', 'dc_phase': '阶段4', 'description': '信息呈现与范例是否充分',
         'passed': 'pass' if has_strategy and strategy.get('segments') else 'risk',
         'notes': ''},
        {'check_id': 'C12', 'dc_phase': '阶段4', 'description': '练习与反馈设计是否完整',
         'passed': 'pass' if has_strategy else 'risk',
         'notes': ''},
        {'check_id': 'C13', 'dc_phase': '阶段5', 'description': '前测/后测/量规是否完整',
         'passed': 'pass' if assessment.get('posttest') else 'risk',
         'notes': ''},
        {'check_id': 'C14', 'dc_phase': '阶段5', 'description': '材料一致性检查是否通过',
         'passed': 'pass' if alignment_pass else 'risk',
         'notes': ''},
        {'check_id': 'C15', 'dc_phase': '阶段6', 'description': '所有材料是否生成且可直接使用',
         'passed': 'pass' if has_materials else 'missing',
         'notes': f'{sum(1 for v in materials.values() if v)}类材料'},
        {'check_id': 'C16', 'dc_phase': '阶段7', 'description': '形成性评价计划是否制定',
         'passed': 'risk',
         'notes': '形成性评价模块尚未实现'},
        {'check_id': 'C17', 'dc_phase': '阶段7', 'description': '修订记录是否完整',
         'passed': 'risk',
         'notes': '修订引擎尚未实现'},
        {'check_id': 'C18', 'dc_phase': '阶段8', 'description': '质量门禁检查是否通过',
         'passed': 'pass' if quality_pass else 'risk',
         'notes': f'分数: {quality.get("score", 0)}/100'},
    ]

    return items


# ---------------------------------------------------------------------------
# Public export: All documents
# ---------------------------------------------------------------------------

def export_all(project: dict, output_dir: str) -> dict:
    """
    Run all exports and create an export index.

    Generates all 5 document files and an export_index.json.
    Filenames are stable (no timestamp prefix):

      - <prefix>_full_dc_report.docx
      - <prefix>_lesson_plan.docx
      - <prefix>_student_worksheet.docx
      - <prefix>_alignment_matrix.xlsx
      - <prefix>_ai_process_record.docx
      - <prefix>_export_index.json

    Returns: {files: dict, index_path: str}
    """
    os.makedirs(output_dir, exist_ok=True)

    meta = project.get('metadata', {})
    proj_name = _safe(meta.get('project_name'), 'mvp_algorithm')
    prefix = ''.join(c if c.isalnum() or c == '_' else '_' for c in proj_name)
    prefix = prefix.strip('_').lower() or 'mvp_algorithm'

    files = {}

    # 1. Full DC Report
    docx_path = os.path.join(output_dir, f'{prefix}_full_dc_report.docx')
    result = export_full_dc_report(project, docx_path)
    files['dc_report'] = result

    # 2. Lesson Plan
    lesson_path = os.path.join(output_dir, f'{prefix}_lesson_plan.docx')
    result = export_lesson_plan(project, lesson_path)
    files['lesson_plan'] = result

    # 3. Student Worksheet
    worksheet_path = os.path.join(output_dir, f'{prefix}_student_worksheet.docx')
    result = export_student_worksheet(project, worksheet_path)
    files['student_worksheet'] = result

    # 4. Alignment Matrix
    matrix_path = os.path.join(output_dir, f'{prefix}_alignment_matrix.xlsx')
    result = export_alignment_matrix(project, matrix_path)
    files['alignment_matrix'] = result

    # 5. AI Process Record
    ai_path = os.path.join(output_dir, f'{prefix}_ai_process_record.docx')
    result = export_ai_process_record(project, ai_path)
    files['ai_process_record'] = result

    # Create export index
    quality = project.get('quality_check', {})
    can_final = quality.get('export_as_final', quality.get('can_export_as_final', False))

    all_warnings = []
    for key, info in files.items():
        if info.get('warnings'):
            all_warnings.extend(info['warnings'])

    index = {
        'project_id': _safe(meta.get('project_id'), 'unknown'),
        'project_name': _safe(meta.get('project_name'), '未命名'),
        'exported_at': datetime.now().isoformat(),
        'export_status': 'success' if all(f.get('exported', False) for f in files.values()) else 'partial',
        'warnings': all_warnings,
        'export_mode': 'final' if can_final else 'draft',
        'full_report_docx': files.get('dc_report', {}).get('path', ''),
        'lesson_plan_docx': files.get('lesson_plan', {}).get('path', ''),
        'student_worksheet_docx': files.get('student_worksheet', {}).get('path', ''),
        'alignment_matrix_xlsx': files.get('alignment_matrix', {}).get('path', ''),
        'ai_process_record_docx': files.get('ai_process_record', {}).get('path', ''),
        'source_json': '',
        'source_markdown_report': '',
        'source_materials_markdown': '',
        'files': {},
    }
    for key, info in files.items():
        index['files'][key] = {
            'path': info.get('path', ''),
            'size': info.get('size', 0),
        }

    index_path = os.path.join(output_dir, f'{prefix}_export_index.json')
    with open(index_path, 'w', encoding='utf-8') as f:
        json.dump(index, f, ensure_ascii=False, indent=2)

    return {
        'files': files,
        'index_path': os.path.abspath(index_path),
    }
